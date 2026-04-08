from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
OUTPUT_PATH = BASE_DIR / "final lab" / "lab_11_428_updated.docx"
IMG_DIR = BASE_DIR / "output" / "lab11_kaggle"
DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"
SOURCE_URL = "https://www.kaggle.com/code/vinayaktiwari28/starter-llm-rag-implementation"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Platform"
    hdr[1].text = "Purpose in the Experiment"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


def add_image(doc, image_path, width_inches):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    para.paragraph_format.space_after = Pt(4)


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(8)


def main():
    overview = IMG_DIR / "lab11_kaggle_overview.png"
    query_section = IMG_DIR / "lab11_kaggle_query_section.png"
    invalid_query = IMG_DIR / "lab11_kaggle_invalid_query.png"

    doc = Document()
    configure_document(doc)

    add_paragraph(doc, "GEN_AI LAB 11", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(
        doc,
        "Implementing a RAG-Based LLM Project Using a Vector Database",
        italic=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_paragraph(doc, f"Name: {NAME}")
    add_paragraph(doc, f"USN: {USN}")
    add_paragraph(doc, f"Date: {DATE_STR}", space_after=10)

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "This report is rebuilt around the Kaggle notebook 'Starter LLM RAG implementation' by Vinayak Tiwari. Instead of using a generic repository-level description, the lab now follows that notebook's actual RAG pipeline, explanations, and visible outputs as the main evidence source.",
    )
    add_paragraph(
        doc,
        f"Source notebook used for this lab update: {SOURCE_URL}",
    )

    add_heading(doc, "Objective")
    for item in [
        "To understand retrieval-augmented generation using a practical notebook implementation.",
        "To study how Llama 2, LangChain, and LlamaIndex can be combined in a RAG pipeline.",
        "To observe how document ingestion, embedding creation, indexing, and query answering work together.",
        "To evaluate real notebook outputs rather than only describing RAG conceptually.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(
        doc,
        [
            ("Kaggle Notebook", "Used as the execution environment and evidence source for the RAG workflow."),
            ("Llama 2 7b-chat-hf", "Used as the LLM backend in the notebook."),
            ("LangChain", "Used for orchestration and integration around the RAG components."),
            ("LlamaIndex", "Used for indexing documents and creating the query engine."),
            ("HuggingFace embeddings", "Used to convert text into vector representations."),
        ],
    )

    add_heading(doc, "Notebook-Based Methodology")
    for item in [
        "Load a CSV-based real-estate dataset and inspect the available fields such as rent, address, beds, baths, and property description.",
        "Convert the property description text into document files so the RAG pipeline can index them.",
        "Create embeddings using a HuggingFace embedding model and build a VectorStoreIndex from the documents.",
        "Define a system prompt that restricts the assistant to real-estate questions and blocks unrelated or objectionable requests.",
        "Run the query engine on sample prompts and inspect the returned answers and source-node traces.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    add_paragraph(
        doc,
        "The notebook explains RAG as a retrieval-first workflow in which relevant context is fetched from indexed documents before answer generation. In this implementation, the source information is a CSV-derived real-estate listing, and the indexed text is written into local document files before vector indexing.",
    )
    add_paragraph(
        doc,
        "The implementation uses LlamaIndex with LangChain support and a HuggingFace embedding model to build the vector-backed retrieval layer. The language model is Llama 2 7b-chat-hf, configured with a strict system prompt that limits the assistant to property-related questions only.",
    )
    add_paragraph(
        doc,
        "A key design point in the notebook is prompt restriction. The assistant is intentionally framed as a smart real-estate chat agent, and the prompt tells it to reject non-real-estate or harmful questions. This means the notebook is not only a basic RAG demo but also a constrained-domain assistant.",
    )

    add_heading(doc, "Results from the Kaggle Notebook")
    add_paragraph(
        doc,
        "The first demonstration query asks the system who it is. The returned response describes the assistant as a property-helping AI and references the indexed property context, which shows that retrieval is connected to the generated answer.",
    )
    add_paragraph(
        doc,
        "The second demonstration query asks about rooms in Sydney. Although the notebook heading suggests this should return a useful recommendation, the captured output visible on the page actually shows a refusal-style answer. This indicates that the prompt constraints may be too strong or imperfectly aligned with the intended domain behavior.",
    )
    add_paragraph(
        doc,
        "The third demonstration query asks for a joke. Here the system correctly refuses because the request is outside the real-estate domain. This is a successful example of prompt-based guardrailing on top of the RAG pipeline.",
    )

    add_heading(doc, "Notebook Screenshot Evidence")
    add_paragraph(
        doc,
        "The following screenshots were captured from the Kaggle notebook page and inserted as direct evidence for Lab 11.",
    )

    if overview.exists():
        add_image(doc, overview, 6.2)
        add_caption(doc, "Figure 1. Kaggle notebook overview showing the notebook title, runtime details, and the table of contents for the RAG implementation.")

    if query_section.exists():
        add_image(doc, query_section, 6.2)
        add_caption(doc, "Figure 2. Kaggle notebook section showing the real-estate query demonstrations and the visible query-engine outputs.")

    if invalid_query.exists():
        add_image(doc, invalid_query, 6.2)
        add_caption(doc, "Figure 3. Kaggle notebook section showing the invalid-query guardrail behavior for an out-of-domain prompt.")

    add_heading(doc, "Observations")
    add_paragraph(
        doc,
        "This notebook is a strong example of an end-to-end beginner RAG workflow because it covers ingestion, embeddings, indexing, querying, and guardrails in one place. It is especially useful for understanding how a domain-specific assistant can be built from a small custom document source.",
    )
    add_paragraph(
        doc,
        "The notebook also reveals an important practical lesson: prompt restrictions can improve safety but may also over-restrict valid domain questions. The visible output for the room-related question suggests that prompt design must be tested carefully, not assumed to work perfectly.",
    )
    add_paragraph(
        doc,
        "Compared with a purely theoretical report, this version is stronger because it is backed by a concrete public notebook, real screenshots, and visible query results from the source implementation.",
    )

    add_heading(doc, "Conclusion")
    add_paragraph(
        doc,
        "Lab 11 has now been fully reoriented around the Kaggle notebook implementation you shared. The report uses that notebook's RAG explanation, model stack, workflow, and captured outputs as the primary basis for the lab.",
    )
    add_paragraph(
        doc,
        "This makes the report more evidence-based and more practical, while also documenting an important real-world insight: retrieval quality, prompt quality, and guardrail design all affect the final usefulness of a RAG assistant.",
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
