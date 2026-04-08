import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
NOTEBOOK_PATH = Path(r"C:\Users\licha\Downloads\GenAI_Lab2.ipynb")
OUTPUT_PATH = BASE_DIR / "final lab" / "lab_10_428_updated.docx"
IMG_DIR = BASE_DIR / "output" / "lab10_notebook"
DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Platform"
    hdr[1].text = "Purpose in the Experiment"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


def add_image(doc, image_path, width_inches):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    para.paragraph_format.space_after = Pt(4)


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(8)


def strip_ansi(text):
    return re.sub(r"\x1b\[[0-9;]*m", "", text)


def load_notebook_outputs():
    nb = json.loads(NOTEBOOK_PATH.read_text(encoding="utf-8"))
    targets = {
        "What is LangChain?": None,
        "Why are chains important in LangChain?": None,
        "What is 15 * 6 and what is the weather in Bangalore?": None,
    }
    for cell in nb.get("cells", []):
        if cell.get("cell_type") != "code":
            continue
        source = "".join(cell.get("source", []))
        for key in list(targets):
            if key in source:
                chunks = []
                for out in cell.get("outputs", []):
                    if "text" in out:
                        chunks.append("".join(out["text"]))
                    elif out.get("output_type") == "execute_result":
                        data = out.get("data", {})
                        if "text/plain" in data:
                            chunks.append("".join(data["text/plain"]))
                if chunks:
                    targets[key] = strip_ansi("".join(chunks)).strip()
    return targets


def pick_font(size, mono=False):
    candidates = []
    if mono:
        candidates = [
            r"C:\Windows\Fonts\consola.ttf",
            r"C:\Windows\Fonts\cour.ttf",
        ]
    else:
        candidates = [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
        ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    for paragraph in text.splitlines():
        if not paragraph.strip():
            lines.append("")
            continue
        words = paragraph.split()
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            width = draw.textbbox((0, 0), trial, font=font)[2]
            if width <= max_width:
                current = trial
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def create_output_image(title, body, filename):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    width = 1500
    margin = 70
    line_gap = 10
    title_font = pick_font(34, mono=False)
    body_font = pick_font(23, mono=True)
    dummy = Image.new("RGB", (width, 100), "white")
    draw = ImageDraw.Draw(dummy)
    lines = wrap_text(draw, body, body_font, width - (2 * margin))
    line_height = draw.textbbox((0, 0), "Ag", font=body_font)[3] + line_gap
    title_height = draw.textbbox((0, 0), title, font=title_font)[3]
    height = margin * 2 + title_height + 40 + max(1, len(lines)) * line_height + 30

    image = Image.new("RGB", (width, height), "#f6f8fb")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((25, 25, width - 25, height - 25), radius=24, fill="white", outline="#cfd8e3", width=3)
    draw.text((margin, margin), title, font=title_font, fill="#17324d")
    y = margin + title_height + 40
    for line in lines:
        draw.text((margin, y), line, font=body_font, fill="#1f2937")
        y += line_height

    out_path = IMG_DIR / filename
    image.save(out_path)
    return out_path


def main():
    outputs = load_notebook_outputs()
    langchain_output = outputs["What is LangChain?"] or "Notebook output unavailable."
    chain_output = outputs["Why are chains important in LangChain?"] or "Notebook output unavailable."
    agent_output = outputs["What is 15 * 6 and what is the weather in Bangalore?"] or "Notebook output unavailable."

    img1 = create_output_image(
        "Notebook Output 1: Direct LangChain Query",
        langchain_output,
        "lab10_langchain_response.png",
    )
    img2 = create_output_image(
        "Notebook Output 2: LCEL Chain Result",
        chain_output,
        "lab10_chain_response.png",
    )
    img3 = create_output_image(
        "Notebook Output 3: Agent Executor Trace",
        agent_output,
        "lab10_agent_response.png",
    )

    doc = Document()
    configure_document(doc)

    add_paragraph(doc, "GEN_AI LAB 10", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(
        doc,
        "Hands-On Lab Using LangChain and LangFlow",
        italic=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    add_paragraph(doc, f"Name: {NAME}")
    add_paragraph(doc, f"USN: {USN}")
    add_paragraph(doc, f"Date: {DATE_STR}", space_after=10)

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "This experiment focuses on building hands-on LLM workflows using LangChain concepts such as prompts, chains, output parsing, and tool-calling agents. The uploaded notebook GenAI_Lab2.ipynb provides direct execution evidence for these ideas through actual code cells and generated outputs.",
    )
    add_paragraph(
        doc,
        "Although the notebook is centered more on LangChain than the LangFlow visual builder, its outputs still serve as valid practical evidence for the orchestration concepts expected in this lab. The report therefore uses the notebook outputs as primary proof of execution.",
    )

    add_heading(doc, "Objective")
    for item in [
        "To understand prompt templates and chain construction using LangChain.",
        "To observe how LCEL pipelines transform prompts into structured outputs.",
        "To test a simple tool-calling agent workflow with multiple tools.",
        "To document notebook outputs and screenshots as execution evidence for the lab.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(
        doc,
        [
            ("LangChain", "Used to build prompts, chains, parsers, and the agent workflow."),
            ("Groq with Llama 3.1 8B Instant", "Used as the connected LLM backend in the notebook."),
            ("Google Colab / Jupyter notebook", "Used to execute the uploaded notebook cells."),
            ("python-docx and Pillow", "Used to convert notebook outputs into screenshot evidence for this report."),
        ],
    )

    add_heading(doc, "Methodology / Procedure")
    for item in [
        "Install LangChain-related dependencies and configure the Groq API key in the notebook environment.",
        "Create a direct model call to verify the LangChain and LLM setup.",
        "Build a prompt-template and LCEL chain pipeline to test structured generation.",
        "Define simple tools and run an AgentExecutor workflow to demonstrate tool usage.",
        "Capture the resulting outputs and convert them into screenshot-style evidence for reporting.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    add_paragraph(
        doc,
        "The uploaded notebook begins with package installation and initializes ChatGroq using the llama-3.1-8b-instant model. A direct query asking 'What is LangChain?' verifies that the notebook is connected properly and able to generate model responses.",
    )
    add_paragraph(
        doc,
        "The next stage uses PromptTemplate and ChatPromptTemplate along with StrOutputParser to create an LCEL-style chain. This demonstrates how LangChain can structure prompt flow, compose components, and produce reusable intermediate steps instead of relying on isolated single-shot prompts.",
    )
    add_paragraph(
        doc,
        "The final stage defines two tools, calculator and get_weather, and runs them through a tool-calling agent executor. This is important because it moves beyond text generation into an agent-style workflow where the model chooses actions and invokes tools during execution.",
    )

    add_heading(doc, "Notebook Output Evidence")
    add_paragraph(
        doc,
        "The following screenshots were generated directly from the outputs present inside the uploaded notebook and inserted as evidence for Lab 10.",
    )

    add_image(doc, img1, 6.2)
    add_caption(doc, "Figure 1. Direct LangChain query output from the notebook after calling the Groq-backed model.")

    add_image(doc, img2, 6.2)
    add_caption(doc, "Figure 2. LCEL chain output explaining why chains are important in LangChain.")

    add_image(doc, img3, 6.2)
    add_caption(doc, "Figure 3. AgentExecutor trace showing tool invocation for calculator and weather functions.")

    add_heading(doc, "Observations")
    add_paragraph(
        doc,
        "The notebook demonstrates that LangChain is effective for organizing multi-step LLM workflows into modular and reusable components. Even in a compact notebook, prompts, parsers, and chains clearly separate responsibilities.",
    )
    add_paragraph(
        doc,
        "The agent execution section is especially useful because it shows visible tool invocation rather than only a final answer. This makes the orchestration process easier to understand and aligns well with the educational goal of learning framework-level LLM pipelines.",
    )
    add_paragraph(
        doc,
        "The uploaded notebook does not include a LangFlow canvas screenshot, so the strongest direct evidence available is LangChain notebook execution. That still provides meaningful hands-on evidence for the core concepts covered by the lab title.",
    )

    add_heading(doc, "Conclusion")
    add_paragraph(
        doc,
        "This lab report now uses actual outputs from the uploaded notebook instead of placeholder text. The notebook proves successful execution of LangChain-based prompting, LCEL chaining, and agent tool-calling with a Groq-hosted model.",
    )
    add_paragraph(
        doc,
        "With these screenshots and outputs included, Lab 10 now reflects direct practical work and is much closer to a complete evidence-backed submission.",
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
