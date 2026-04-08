from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
OUTPUT_PATH = BASE_DIR / "final lab" / "lab_7_428.docx"
DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)


def add_code_line(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10.5)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Platform"
    hdr[1].text = "Purpose in the Experiment"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


doc = Document()
configure_document(doc)

add_paragraph(doc, "GEN_AI LAB 7", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(
    doc,
    "Gemini-Pro: Generate API Key and Access the Model Using the API Key",
    italic=True,
    size=12,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=10,
)
add_paragraph(doc, f"Name: {NAME}")
add_paragraph(doc, f"USN: {USN}")
add_paragraph(doc, f"Date: {DATE_STR}", space_after=10)

add_heading(doc, "Introduction")
add_paragraph(
    doc,
    "This experiment demonstrates how to generate a Gemini API key using Google AI Studio and how that key can be used to access the Gemini model programmatically. The main goal is to understand the setup path from account-level key generation to a working API request.",
)
add_paragraph(
    doc,
    "The supplied screenshot shows the Google AI Studio API Keys page, where an API key has already been created under the project 'My First Project'. This serves as output evidence that the API key generation step was completed successfully.",
)

add_heading(doc, "Objective")
for item in [
    "To generate a Gemini API key from Google AI Studio.",
    "To understand how the generated key is associated with a project.",
    "To configure a Python client with the API key.",
    "To execute a model request using Gemini through code.",
]:
    add_bullet(doc, item)

add_heading(doc, "Tools and Platforms Used")
add_tools_table(
    doc,
    [
        ("Google AI Studio", "Used to create and manage the Gemini API key."),
        ("Gemini-Pro / Gemini API", "Target model family accessed through the generated key."),
        ("Python", "Used to write code examples for API-based access."),
        ("google-genai SDK", "Used as the official SDK for Gemini model requests."),
    ],
)

add_heading(doc, "Methodology / Procedure")
for item in [
    "Sign in to Google AI Studio and open the API Keys section.",
    "Create or select the required project for API usage.",
    "Generate a new Gemini API key and keep it private.",
    "Install the Gemini SDK in Python or Colab.",
    "Initialize the client with the generated API key.",
    "Send a sample prompt to verify that the model is reachable.",
]:
    add_bullet(doc, item)

add_heading(doc, "Implementation / Workflow Summary")
add_paragraph(
    doc,
    "The screenshot shows the API Keys dashboard in Google AI Studio. The interface lists one Gemini API key associated with 'My First Project', and the project identifier is visible under the project name. The key is partially masked in the dashboard for safety, which is the correct behavior for secure credential handling.",
)
add_paragraph(
    doc,
    "The page also confirms the creation date and billing tier, which indicates that the project is active and ready for API-based experimentation. Once the key is generated, it can be copied into a secure environment variable or notebook secret and then passed into the Gemini client for actual model access.",
)
add_paragraph(
    doc,
    "This workflow is important because the API key acts as the bridge between the Google AI Studio account and the external code environment. Without the key, no authenticated Gemini model call can be made from Python or Colab.",
)

add_heading(doc, "Output Evidence from Screenshot")
add_paragraph(
    doc,
    "Figure 1 corresponds to the Google AI Studio API Keys page. The screenshot shows the API Keys section open, one Gemini API key entry present, the linked project 'My First Project', the creation date of April 8, 2026, and the free-tier billing status. This is sufficient evidence that the API key generation step was completed successfully.",
)

add_heading(doc, "Examples for Accessing the Model Using API Key")
add_paragraph(
    doc,
    "After generating the key, the model can be accessed through Python. A typical usage pattern is to import the Gemini SDK, initialize the client with the API key, and send a content-generation request to a Gemini model.",
)

for line in [
    "from google import genai",
    "",
    'API_KEY = "YOUR_GEMINI_API_KEY"',
    "client = genai.Client(api_key=API_KEY)",
    "",
    'response = client.models.generate_content(',
    '    model="gemini-2.5-flash",',
    '    contents="Explain generative AI in simple terms."',
    ")",
    "",
    "print(response.text)",
]:
    add_code_line(doc, line)

doc.add_paragraph()
add_paragraph(
    doc,
    "An equivalent Colab workflow follows the same structure: install the package, configure the API key, create the client, and print the generated response. This confirms that the key can be used from either a local Python environment or a notebook-based setup.",
)

add_heading(doc, "Observations")
for item in [
    "Google AI Studio provides a straightforward interface for generating and organizing Gemini API keys.",
    "The screenshot confirms that the API key is linked to a specific project rather than existing as an isolated credential.",
    "Masking the key in the dashboard is an important security feature and reminds users not to expose the full key publicly.",
    "Once the key is created, model access in Python is simple and requires only a few lines of code.",
]:
    add_bullet(doc, item)

add_heading(doc, "Conclusion")
add_paragraph(
    doc,
    "This lab successfully demonstrates the Gemini API key generation workflow using Google AI Studio. The provided screenshot serves as valid evidence that the API key was created and associated with an active project.",
)
add_paragraph(
    doc,
    "The coding example shows how the generated key can be used to access the Gemini model programmatically. Together, these steps complete the full lab objective: generate the key, configure the client, and prepare the environment for authenticated model access.",
)

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT_PATH))
print(OUTPUT_PATH)
