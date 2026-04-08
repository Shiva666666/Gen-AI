import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
FINAL_LAB_DIR = BASE_DIR / "final lab"
LAB12_DATA = json.loads((BASE_DIR / "output" / "lab12_notebook" / "lab12_results.json").read_text(encoding="utf-8"))
LAB13_DATA = json.loads((BASE_DIR / "output" / "lab13_notebook" / "lab13_results.json").read_text(encoding="utf-8"))
LAB14_DATA = json.loads((BASE_DIR / "output" / "lab14_notebook" / "lab14_results.json").read_text(encoding="utf-8"))

DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Platform"
    hdr[1].text = "Purpose in the Experiment"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


def add_image(doc, image_path, width_inches):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    para.paragraph_format.space_after = Pt(4)


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(8)


def add_title_block(doc, lab_no, title):
    add_paragraph(doc, f"GEN_AI LAB {lab_no}", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, title, italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc, f"Name: {NAME}")
    add_paragraph(doc, f"USN: {USN}")
    add_paragraph(doc, f"Date: {DATE_STR}", space_after=10)


def build_lab12():
    doc = Document()
    configure_document(doc)
    add_title_block(doc, 12, "Building a Simple LLM Agent Using the Phi Data Framework")

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "This lab was implemented using Agno, which is the current successor to the Phi Data framework. The agent was executed locally with Ollama and the free `qwen2.5:3b` model, so the report is based on real runtime outputs rather than a manual-only description.",
    )

    add_heading(doc, "Objective")
    for item in [
        "To build a simple role-based LLM agent using the modern Agno framework.",
        "To configure a local model backend through Ollama.",
        "To observe how the same agent responds consistently across multiple prompts.",
        "To document the execution evidence with notebook-style output captures.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(
        doc,
        [
            ("Agno", "Used as the agent framework; documented as the current successor to Phi Data."),
            ("Ollama", "Used to run the local open-source model backend."),
            ("qwen2.5:3b", "Free local model used by the agent."),
            ("Jupyter Notebook", "Used to store the experiment workflow and outputs."),
        ],
    )

    add_heading(doc, "Methodology / Procedure")
    for item in [
        "Install Agno and confirm local Ollama access.",
        "Define a study-assistant agent with role, instructions, and concise-response behavior.",
        "Run three prompts covering concept explanation, applications, and viva preparation.",
        "Save the notebook and capture the resulting outputs as evidence images.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    add_paragraph(
        doc,
        "The notebook initializes an Agno agent called Campus Study Assistant and attaches the `qwen2.5:3b` model through the Ollama provider. The instructions guide the agent to answer in a practical, student-friendly style suitable for a Gen AI lab record or viva preparation.",
    )
    add_paragraph(
        doc,
        f"The executed notebook file is stored at: {LAB12_DATA['notebook']}",
    )

    add_heading(doc, "Observed Agent Prompts")
    for record in LAB12_DATA["records"]:
        add_bullet(doc, f"{record['prompt']}")

    add_heading(doc, "Screenshot Evidence")
    for idx, image in enumerate(LAB12_DATA["images"], start=1):
        add_image(doc, Path(image), 6.2)
        add_caption(doc, f"Figure {idx}. Local Agno agent response captured from the Lab 12 execution workflow.")

    add_heading(doc, "Observations")
    add_paragraph(
        doc,
        "The agent responded coherently across all three prompts and preserved the intended helpful study-assistant tone. This shows that the agent abstraction is working as more than a one-off prompt; the role and instructions remain active across repeated interactions.",
    )
    add_paragraph(
        doc,
        "The local Ollama setup also makes the experiment reproducible without paid APIs, which is useful for future lab revisions or viva demonstrations.",
    )

    add_heading(doc, "Result / Conclusion")
    add_paragraph(
        doc,
        "The lab was successfully completed using a free local stack based on Agno and Ollama. The notebook and screenshots confirm that a simple LLM agent was created, configured, and tested successfully.",
    )

    out_path = FINAL_LAB_DIR / "lab_12_428.docx"
    doc.save(str(out_path))


def build_lab13():
    doc = Document()
    configure_document(doc)
    add_title_block(doc, 13, "Using LLMs for Code Generation and Bug Detection in Software Development")

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "This lab was implemented with the free local `qwen2.5-coder:3b` model running through Ollama. The workflow demonstrates both code generation and bug detection, followed by verification of the corrected Python code.",
    )

    add_heading(doc, "Objective")
    for item in [
        "To generate Python code from a software requirement prompt.",
        "To identify a logical bug in faulty source code using an LLM.",
        "To obtain a corrected version of the program and run it locally.",
        "To capture notebook-style evidence for all major outputs.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(
        doc,
        [
            ("Ollama", "Used to run the local coding model."),
            ("qwen2.5-coder:3b", "Free code-generation and debugging model used in the experiment."),
            ("Python", "Used to execute the corrected code for verification."),
            ("Jupyter Notebook", "Used to organize the prompt-response workflow."),
        ],
    )

    add_heading(doc, "Methodology / Procedure")
    for item in [
        "Ask the coding model to generate a StudentRecordManager class from a textual requirement.",
        "Provide a faulty average-calculation program and request a bug explanation plus fix.",
        "Extract the corrected code and run it with Python for verification.",
        "Store the outputs in the notebook and export them as screenshot evidence.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    add_paragraph(
        doc,
        "The generated program correctly introduced a `StudentRecordManager` class with methods for adding students, computing class average, and listing passed students. For the debugging task, the model identified the off-by-one loop error in the faulty averaging function and proposed a correct replacement.",
    )
    add_paragraph(
        doc,
        f"The corrected program was executed locally and produced the output `{LAB13_DATA['verification_output']}`, confirming that the fix works in practice.",
    )
    add_paragraph(
        doc,
        f"The executed notebook file is stored at: {LAB13_DATA['notebook']}",
    )

    add_heading(doc, "Screenshot Evidence")
    for idx, image in enumerate(LAB13_DATA["images"], start=1):
        add_image(doc, Path(image), 6.2)
        add_caption(doc, f"Figure {idx}. Captured output from the Lab 13 code-generation and bug-fixing workflow.")

    add_heading(doc, "Observations")
    add_paragraph(
        doc,
        "The model was effective for both requirement-to-code generation and reasoning about a runtime bug. In the verification phase, the final corrected function was even simplified further by replacing manual looping with Python's built-in `sum()` operation.",
    )
    add_paragraph(
        doc,
        "This experiment also shows why human validation remains important: the generated answer must still be executed and reviewed before being accepted as correct software behavior.",
    )

    add_heading(doc, "Result / Conclusion")
    add_paragraph(
        doc,
        "The lab was completed successfully using a free local coding model. Real outputs were generated for code creation, bug analysis, and corrected-code verification, and those outputs are embedded in this report as evidence.",
    )

    out_path = FINAL_LAB_DIR / "lab_13_428.docx"
    doc.save(str(out_path))


def build_lab14():
    doc = Document()
    configure_document(doc)
    add_title_block(doc, 14, "Using Multimodal LLMs for Visual Question Answering")

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "This lab was implemented with Hugging Face Transformers using the `dandelin/vilt-b32-finetuned-vqa` model. A local image was used for the experiment, and three visual questions were answered with confidence scores recorded from the model output.",
    )

    add_heading(doc, "Objective")
    for item in [
        "To perform Visual Question Answering with a free pretrained multimodal model.",
        "To load an image locally and submit multiple text questions about it.",
        "To record the top predicted answer and confidence for each question.",
        "To document the image and answers with screenshot evidence.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(
        doc,
        [
            ("Hugging Face Transformers", "Used to load the visual question answering pipeline."),
            ("dandelin/vilt-b32-finetuned-vqa", "Free VQA model used for the experiment."),
            ("Pillow", "Used to open the local image."),
            ("Jupyter Notebook", "Used to run the VQA workflow and save the outputs."),
        ],
    )

    add_heading(doc, "Methodology / Procedure")
    for item in [
        "Load the local sample image into Python.",
        "Initialize the VQA pipeline with the ViLT model.",
        "Ask three questions about the image and record the top answer with its score.",
        "Save the executed notebook and create screenshot-style output evidence.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    add_paragraph(
        doc,
        "The selected image is a COCO sample showing two cats resting on a pink couch or blanket. The model was asked about object count, scene context, and color. The top answers returned by the model were `2`, `bed`, and `pink`, each with strong confidence.",
    )
    add_paragraph(
        doc,
        f"The executed notebook file is stored at: {LAB14_DATA['notebook']}",
    )

    add_heading(doc, "Screenshot Evidence")
    add_image(doc, Path(LAB14_DATA["image_path"]), 5.8)
    add_caption(doc, "Figure 1. Local image used for the Visual Question Answering experiment.")
    add_image(doc, Path(LAB14_DATA["result_image"]), 6.2)
    add_caption(doc, "Figure 2. Captured VQA results showing the questions, answers, and scores.")

    add_heading(doc, "Observations")
    add_paragraph(
        doc,
        "The model correctly counted the cats and identified the dominant pink seating surface with high confidence. The second answer, `bed`, shows that VQA models can still produce scene labels that are semantically close but not always exact, which is a useful limitation to note in a lab report.",
    )
    add_paragraph(
        doc,
        "Overall, the experiment demonstrates that multimodal models can combine image understanding with natural-language questions effectively without depending on paid APIs.",
    )

    add_heading(doc, "Result / Conclusion")
    add_paragraph(
        doc,
        "The lab was completed successfully using a free Hugging Face VQA model. The report now includes both the actual image and the captured answer summary from the executed notebook workflow.",
    )

    out_path = FINAL_LAB_DIR / "lab_14_428.docx"
    doc.save(str(out_path))


def build_lab15():
    doc = Document()
    configure_document(doc)
    add_title_block(doc, 15, "Course Project Using a Free Local VQA Model with a Gradio Frontend")

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "For the course-project lab, a complete local web application was built with Gradio and Hugging Face Transformers. Instead of using Llama or Gemini, the project uses the same free ViLT visual question answering model from Lab 14 and exposes it through an interactive frontend.",
    )

    add_heading(doc, "Objective")
    for item in [
        "To build a small but complete AI project with a user-facing frontend.",
        "To reuse the VQA model inside a Gradio application.",
        "To demonstrate image upload, question input, answer generation, and top-k predictions.",
        "To capture a real browser screenshot of the running application.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(
        doc,
        [
            ("Gradio", "Used to build the frontend interface."),
            ("Hugging Face Transformers", "Used to run the VQA model in the backend."),
            ("dandelin/vilt-b32-finetuned-vqa", "Free local multimodal model used in the app."),
            ("Microsoft Edge (headless screenshot)", "Used to capture the running application as evidence."),
        ],
    )

    add_heading(doc, "Methodology / Procedure")
    for item in [
        "Create a Gradio interface with image input, question input, answer output, and a prediction table.",
        "Load the ViLT VQA model once when the app starts.",
        "Preload the sample cat image and default question so the app demonstrates a ready-to-use prediction.",
        "Launch the app locally and capture a browser screenshot showing the final result.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    add_paragraph(
        doc,
        "The application file `output\\gradio\\lab_15_vqa_app.py` builds a compact image-question-answering assistant. The left side shows the input image, and the right side provides a textbox for the question, the top predicted answer, and a table of the top three predictions with scores.",
    )
    add_paragraph(
        doc,
        "The default question used in the app is `How many cats are there?`, and the displayed answer is `2` with a confidence of approximately `0.8799`. This makes the project a valid end-to-end Gen AI mini application with a usable local frontend.",
    )

    add_heading(doc, "Screenshot Evidence")
    add_image(doc, BASE_DIR / "output" / "playwright" / "lab15_app_ready.png", 6.3)
    add_caption(doc, "Figure 1. Running Gradio frontend showing the image, question, top answer, and top-3 predictions.")
    add_image(doc, BASE_DIR / "output" / "playwright" / "lab14_sample_cats.jpg", 5.8)
    add_caption(doc, "Figure 2. Input image used by the local course-project application.")

    add_heading(doc, "Observations")
    add_paragraph(
        doc,
        "This project demonstrates how a notebook experiment can be promoted into a small deployable interface. The frontend is lightweight, but it still covers the core user journey: choose an image, ask a question, and inspect the model's response.",
    )
    add_paragraph(
        doc,
        "Using a free local model keeps the project practical for student systems and avoids dependency on commercial APIs while still satisfying the course-project requirement.",
    )

    add_heading(doc, "Result / Conclusion")
    add_paragraph(
        doc,
        "Lab 15 was successfully implemented as a free local Gradio-based VQA project. The report now contains a real application screenshot with visible predictions instead of placeholder text, making it suitable as execution evidence for the final lab record.",
    )

    out_path = FINAL_LAB_DIR / "lab_15_428.docx"
    doc.save(str(out_path))


def main():
    FINAL_LAB_DIR.mkdir(parents=True, exist_ok=True)
    build_lab12()
    build_lab13()
    build_lab14()
    build_lab15()
    print("Updated lab 12 to lab 15 reports.")


if __name__ == "__main__":
    main()
