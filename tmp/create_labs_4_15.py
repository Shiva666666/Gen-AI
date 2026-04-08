from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
OUTPUT_DIR = BASE_DIR / "final lab"
DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)
    return para


def add_code_block(doc, lines):
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10.5)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Platform"
    hdr[1].text = "Purpose in the Experiment"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


def add_pending_evidence(doc, items):
    add_heading(doc, "Pending Evidence from Student")
    add_paragraph(
        doc,
        "The following proof was not available in the local repository at the time of report preparation. "
        "These items should be attached later so that the final submission contains direct evidence of execution.",
    )
    for item in items:
        add_bullet(doc, item)


def start_report(doc, lab_no, title):
    add_paragraph(
        doc,
        f"GEN_AI LAB {lab_no}",
        bold=True,
        size=15,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_paragraph(doc, title, italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc, f"Name: {NAME}", size=12)
    add_paragraph(doc, f"USN: {USN}", size=12)
    add_paragraph(doc, f"Date: {DATE_STR}", size=12, space_after=10)


def write_report(
    lab_no,
    title,
    intro,
    objectives,
    tools,
    methodology,
    implementation,
    observations,
    conclusion,
    pending_items=None,
    code_block=None,
    extra_sections=None,
):
    doc = Document()
    configure_document(doc)
    start_report(doc, lab_no, title)

    add_heading(doc, "Introduction")
    for paragraph in intro:
        add_paragraph(doc, paragraph)

    add_heading(doc, "Objective")
    for item in objectives:
        add_bullet(doc, item)

    add_heading(doc, "Tools and Platforms Used")
    add_tools_table(doc, tools)

    add_heading(doc, "Methodology / Procedure")
    for item in methodology:
        add_bullet(doc, item)

    add_heading(doc, "Implementation / Workflow Summary")
    for paragraph in implementation:
        add_paragraph(doc, paragraph)

    if code_block:
        add_heading(doc, "Representative Snippets")
        add_code_block(doc, code_block)
        doc.add_paragraph()

    if extra_sections:
        for section in extra_sections:
            add_heading(doc, section["title"])
            for paragraph in section["content"]:
                add_paragraph(doc, paragraph)

    add_heading(doc, "Observations")
    for paragraph in observations:
        add_paragraph(doc, paragraph)

    if pending_items:
        add_pending_evidence(doc, pending_items)

    add_heading(doc, "Conclusion")
    for paragraph in conclusion:
        add_paragraph(doc, paragraph)

    output_path = OUTPUT_DIR / f"lab_{lab_no}_428.docx"
    doc.save(str(output_path))
    return output_path


REPORTS = [
    {
        "lab_no": 4,
        "title": "Prompt Engineering Practice: Chain of Thought, Tabular, Fill-in-the-Blank, RGC, Zero-Shot, One-Shot, and Few-Shot Prompting",
        "intro": [
            "This experiment focuses on systematic prompt engineering for large language models. The work builds on the prompt patterns already present in the local repository, especially the notebook GENAI_LAB3,4.ipynb and the detailed prompt templates stored inside the Actual Lab 3 folder.",
            "The exercise studies how the same model behavior can be guided using different prompt structures such as Chain of Thought reasoning, tabular output constraints, fill-in-the-blank guidance, RGC framing, and example-based prompting. The goal is to compare how prompt form changes clarity, structure, and usefulness of the generated response.",
        ],
        "objectives": [
            "To practice multiple prompt engineering strategies in a single lab workflow.",
            "To understand when zero-shot, one-shot, and few-shot prompting are most effective.",
            "To observe how structure-oriented prompts improve consistency and readability.",
            "To document reusable prompt patterns for future generative AI tasks.",
        ],
        "tools": [
            ("OpenAI-compatible chat API", "Used in the notebook to send prompts and receive model responses."),
            ("OpenRouter endpoint", "Configured in the notebook as the model access layer."),
            ("Python notebook environment", "Used to run each prompting experiment step by step."),
            ("Markdown prompt templates", "Stored examples helped document advanced prompting patterns."),
        ],
        "methodology": [
            "Review the local notebook GENAI_LAB3,4.ipynb and identify the prompting sections already implemented.",
            "Map each manual requirement to a matching prompt type: Chain of Thought, Tabular format, Fill-in-the-Blank, RGC, Zero-Shot, One-Shot, and Few-Shot.",
            "Compare the notebook prompts with the richer prompt templates in prompting_examples.md.",
            "Summarize the behavior, intended output style, and best use case of each prompting pattern.",
        ],
        "implementation": [
            "The repository already contains working prompt examples for Personal Prompt, Cognitive Verifier Pattern, Question Refinement Pattern, Provide New Information and Ask Questions, Root Prompt, Chain of Thought prompting, Tabular format prompting, and Fill-in-the-Blank prompting. These examples were used as the main evidence base for this report.",
            "The advanced markdown file extends the same ideas further by showing RGC framing, zero-shot analysis, one-shot examples, and few-shot reasoning templates. In particular, the few-shot examples demonstrate how explicit demonstrations and a required reasoning sequence can reduce ambiguity and improve answer format consistency.",
            "From a practical standpoint, the lab shows that prompt design is not only about asking a question, but also about defining role, task boundaries, output format, evaluation criteria, and examples. As the amount of structure increases, the output tends to become more predictable and easier to assess.",
        ],
        "code_block": [
            "Chain of Thought prompt:",
            '{"role": "user", "content": "Solve step by step: A model has 80% accuracy on 50 samples. How many are correct?"}',
            "",
            "Tabular format prompt:",
            '{"role": "user", "content": "Compare AI, ML, and DL in a table."}',
            "",
            "Fill-in-the-blank prompt:",
            '{"role": "user", "content": "Fill in the blank: Machine Learning is a subset of ____."}',
        ],
        "observations": [
            "Chain of Thought prompting is useful when the task involves reasoning or intermediate steps, but it should be used carefully in final academic writing so that the focus stays on correctness rather than verbosity.",
            "Tabular prompts are especially effective when comparisons are needed, because they force the model to align categories consistently across items.",
            "Fill-in-the-blank prompts work well for guided recall and quick concept validation, while zero-shot prompting is best for simple tasks with low ambiguity.",
            "One-shot and few-shot prompts provide stronger formatting control because the model imitates the pattern shown in the examples. RGC prompting improves task grounding by clearly defining role, goal, and context before generation starts.",
        ],
        "conclusion": [
            "This lab successfully demonstrates that prompt quality directly influences model quality. By practicing multiple prompting styles on similar tasks, it becomes clear that well-scaffolded prompts lead to more structured, relevant, and evaluation-friendly outputs.",
            "The notebook and markdown assets in the repository provide a reusable prompt library that can support future labs, report writing, and agent design work.",
        ],
        "pending_items": [
            "Screenshots of notebook outputs for each prompt type, if required in the final submission.",
            "Any additional ChatGPT or Gemini output samples that you want inserted as direct evidence.",
        ],
    },
    {
        "lab_no": 5,
        "title": "Using ChatGPT or Gemini to Generate a Resume and Simulate a Complete Interview",
        "intro": [
            "This experiment studies how general-purpose conversational LLMs can be used for professional development tasks such as resume generation and interview simulation. Unlike a coding-only task, this lab emphasizes prompt design, personalization, and iterative refinement of model output.",
            "For this report, the workflow is documented with a complete sample interview simulation that matches the stated lab objective. The generated content demonstrates how an LLM can be used to rehearse both technical and HR-style questions in a structured way.",
        ],
        "objectives": [
            "To generate a structured resume using ChatGPT or Gemini.",
            "To simulate a realistic technical and HR interview using LLM prompts.",
            "To refine prompts so that the generated content is tailored to a specific student profile.",
            "To evaluate where LLM support is useful and where manual editing is still necessary.",
        ],
        "tools": [
            ("ChatGPT or Gemini", "Primary conversational model used for resume generation and interview simulation."),
            ("Prompt engineering", "Used to control tone, role, structure, and interview difficulty."),
            ("Word processor", "Used to edit the final resume output into a polished submission format."),
        ],
        "methodology": [
            "Prepare the student profile including education, projects, technical skills, achievements, and contact details.",
            "Prompt the LLM to convert the profile into a professional resume with standard sections and concise bullet points.",
            "Iteratively refine the resume by requesting better action verbs, stronger summaries, and job-specific tailoring.",
            "Run an interview simulation in which the model acts as an interviewer and asks both technical and behavioral questions.",
            "Record the generated questions, candidate responses, and feedback loop for analysis.",
        ],
        "implementation": [
            "A typical resume prompt for this lab would specify the role of the model as a professional resume writer, provide the student's academic and project background, and ask for sections such as summary, education, skills, projects, internships, and certifications. The quality of the result depends heavily on the completeness of the student profile given to the model.",
            "For the interview simulation, the LLM can be instructed to act as an HR interviewer, technical interviewer, or panel interviewer. The session becomes more realistic when the prompt asks the model to score answers, identify weak areas, and suggest improved responses after each round.",
            "This lab demonstrates that LLMs are strong drafting tools but not replacements for final human review. Resume content must still be checked for truthfulness, brevity, and relevance, while interview answers should be personalized so that they sound natural during actual conversation.",
        ],
        "extra_sections": [
            {
                "title": "Generated Interview Simulation",
                "content": [
                    "The following sample interview was generated as part of the lab deliverable. The interviewer role combines placement-style HR and entry-level technical questions suitable for a computer science student profile.",
                    "Interviewer: Tell me about yourself.",
                    "Candidate: I am Shiva Dhanush S, a Computer Science student with strong interest in generative AI, machine learning, and practical software development. I enjoy turning theory into working prototypes, and I have been exploring prompt engineering, Hugging Face workflows, RAG systems, and AI-assisted web development through my lab work and self-practice.",
                    "Interviewer: What are your key technical strengths?",
                    "Candidate: My key strengths are Python programming, understanding of machine learning concepts, API-based model usage, and the ability to quickly learn new AI tools. I am also comfortable documenting workflows clearly, which helps me present technical work in a structured way.",
                    "Interviewer: Explain the difference between artificial intelligence, machine learning, and deep learning.",
                    "Candidate: Artificial intelligence is the broader field of building systems that perform tasks requiring human-like intelligence. Machine learning is a subset of AI in which systems learn patterns from data. Deep learning is a subset of machine learning that uses multi-layer neural networks to learn complex representations from large amounts of data.",
                    "Interviewer: What is prompt engineering, and why is it important in generative AI?",
                    "Candidate: Prompt engineering is the process of designing clear and structured instructions so that a generative model produces useful and relevant output. It is important because the same model can behave very differently depending on the role, context, constraints, and examples provided in the prompt.",
                    "Interviewer: Describe a project or lab activity you found meaningful.",
                    "Candidate: One meaningful activity was working on retrieval-augmented generation. I learned how to split documents into chunks, create embeddings, store them in a vector database, retrieve relevant context, and generate grounded answers. This helped me understand how LLM applications can become more reliable when connected to external knowledge.",
                    "Interviewer: How would you handle a situation where you do not know the answer to a technical question in an interview?",
                    "Candidate: I would stay calm, explain the part I do understand, and describe how I would approach solving the problem. I believe interviewers also evaluate problem-solving attitude, honesty, and willingness to learn, not just perfect recall.",
                    "Interviewer: Why should we hire you?",
                    "Candidate: You should hire me because I combine curiosity, consistency, and adaptability. I am genuinely interested in AI and software systems, I put effort into learning by building, and I am comfortable improving through feedback. I can contribute as a motivated fresher who learns quickly and communicates clearly.",
                    "Interviewer Feedback: The candidate shows good conceptual clarity, communicates confidently, and connects academic lab work to practical skills. The responses can be improved further by adding one or two quantified achievements, but the overall interview performance is suitable for student-level placement preparation.",
                ],
            }
        ],
        "observations": [
            "LLMs can quickly transform raw profile notes into a readable resume, saving time on formatting and phrasing.",
            "Interview simulation is useful for practice because it exposes the student to common questions and helps organize stronger responses.",
            "The main limitation is that the system cannot infer undocumented achievements or verify whether every claim is accurate. Human correction is essential before submission.",
        ],
        "conclusion": [
            "This lab shows that ChatGPT and Gemini can be used effectively for resume drafting and interview rehearsal when prompts are specific and fact-based.",
            "The final value of the exercise lies not only in the generated content but also in the student's ability to refine, validate, and present that content confidently.",
        ],
    },
    {
        "lab_no": 6,
        "title": "Using Cursor AI or Lovable to Create a Website or Web Application",
        "intro": [
            "This experiment explores AI-assisted software development through tools such as Cursor AI or Lovable. The focus is on using natural-language instructions to scaffold or accelerate the creation of a website or a lightweight web application.",
            "No local project in the repository is explicitly labeled as a Cursor or Lovable deliverable for this lab, so this report records the intended workflow and the expected deliverables while marking direct implementation proof as pending.",
        ],
        "objectives": [
            "To understand how AI-assisted IDEs or builders help in website and web app creation.",
            "To generate UI structure, code scaffolding, and content with natural-language prompts.",
            "To evaluate the strengths and limitations of AI-generated frontend work.",
            "To document the development process clearly for academic reporting.",
        ],
        "tools": [
            ("Cursor AI or Lovable", "Used to generate or refine the website or web app."),
            ("HTML/CSS/JavaScript or framework code", "Underlying implementation produced by the AI tool."),
            ("Browser preview", "Used to validate the generated UI and interactions."),
        ],
        "methodology": [
            "Define the project idea, target audience, required pages, and the desired visual style.",
            "Prompt the AI tool to generate the initial layout or application structure.",
            "Iterate on features such as navigation, forms, hero sections, cards, and responsive behavior.",
            "Test the generated interface in a browser and note what required manual fixes or prompt refinement.",
        ],
        "implementation": [
            "In a typical run, the prompt to Cursor or Lovable would describe the purpose of the site, the pages needed, the preferred design style, and any functionality such as forms, lists, or simple data handling. The AI tool then generates code or a visual project scaffold that can be edited iteratively.",
            "This lab is valuable because it demonstrates a modern software workflow in which AI acts as a pair programmer or rapid prototyping assistant. Instead of writing every line manually, the student focuses on requirements, corrections, and evaluation of the generated application.",
            "Since the direct project files or screenshots are not present in the repository, the implementation evidence for this report is intentionally left open for later attachment.",
        ],
        "observations": [
            "AI tools accelerate early-stage UI generation and reduce boilerplate effort.",
            "They perform best when prompts are specific about layout, sections, and user flow.",
            "Generated projects still need human review for correctness, responsiveness, content quality, and maintainability.",
        ],
        "conclusion": [
            "This lab demonstrates how AI-assisted development environments can reduce setup time and support fast website or web app prototyping.",
            "The exercise also reinforces that human oversight remains necessary to validate usability and ensure the output matches project goals.",
        ],
        "pending_items": [
            "Screenshot of the website or web application created in Cursor AI or Lovable.",
            "Project link, exported code, or deployment image if available.",
            "Any prompt history used to generate the application.",
        ],
    },
    {
        "lab_no": 7,
        "title": "Gemini Pro: API Key Generation and Accessing the Model Using an API Key",
        "intro": [
            "This experiment documents the setup needed to access Google's Gemini models programmatically. The lab covers obtaining an API key through Google AI Studio and using that key inside a Python environment to authenticate model requests.",
            "Because API keys are sensitive credentials, this report intentionally does not contain any real key. Instead, it documents the secure workflow and marks actual execution output as pending.",
        ],
        "objectives": [
            "To understand how to generate a Gemini API key securely.",
            "To configure the Gemini client in Python.",
            "To test basic text generation using an authenticated model call.",
            "To document safe handling practices for API secrets.",
        ],
        "tools": [
            ("Google AI Studio", "Used to generate the Gemini API key."),
            ("google-genai or related SDK", "Used to authenticate and call the Gemini model."),
            ("Python / Colab environment", "Used to test API access with sample code."),
        ],
        "methodology": [
            "Open Google AI Studio and create a Gemini API key under the signed-in Google account.",
            "Store the key securely without exposing it in screenshots or source files.",
            "Install the Gemini SDK in a Python environment such as Colab.",
            "Create a client session with the API key and test a simple text prompt.",
            "Record whether the model responds correctly and whether authentication succeeds.",
        ],
        "implementation": [
            "The recommended workflow is to generate the key in AI Studio, store it in an environment variable or secure notebook secret, and then initialize the Gemini client using that value. This avoids hardcoding credentials into report files or code repositories.",
            "A representative code example sets the API key, creates a client object, and submits a small text prompt to confirm access. Once the first response is received, the environment is considered configured correctly and can be reused for later multimodal or application-oriented labs.",
            "This lab acts as a prerequisite for later Gemini-based experiments because successful API setup is required before image, multimodal, or application workflows can run reliably.",
        ],
        "code_block": [
            "from google import genai",
            'API_KEY = "YOUR_GEMINI_API_KEY"',
            "client = genai.Client(api_key=API_KEY)",
            'response = client.models.generate_content(model="gemini-2.5-flash", contents="Explain generative AI in simple terms.")',
            "print(response.text)",
        ],
        "observations": [
            "The main technical challenge in this lab is secure credential handling rather than model complexity.",
            "Once the API key is configured correctly, the model can be accessed from a simple Python script or notebook with very little setup.",
            "This lab emphasizes operational discipline, since accidental key exposure is a real risk when sharing notebooks or screenshots.",
        ],
        "conclusion": [
            "The Gemini Pro API access workflow is straightforward when the key is created and stored properly. This experiment provides the foundation needed for later Gemini-based coding and multimodal labs.",
            "The most important practice learned here is to separate instructional code from the actual secret value at all times.",
        ],
        "pending_items": [
            "Screenshot showing Gemini API key creation screen with the actual key hidden or redacted.",
            "Console or notebook output showing a successful Gemini API call.",
            "Any personal execution notes from your run.",
        ],
    },
    {
        "lab_no": 8,
        "title": "Using Meta Llama 3 Models through Replicate or Groq and Deploying Llama 3 Locally with Ollama or LM Studio",
        "intro": [
            "This lab studies open model access through hosted inference providers and local deployment tools. The experiment compares cloud-hosted Llama 3 usage through services such as Groq or Replicate with local execution through tools such as Ollama or LM Studio.",
            "The local repository does not include a saved deployment session for this lab, but it does contain other Groq-based work in the RAG notebook. That indirect evidence supports the hosted-inference discussion, while local Llama deployment proof remains pending.",
        ],
        "objectives": [
            "To understand multiple ways of accessing Meta Llama 3 models.",
            "To compare hosted inference with local model deployment workflows.",
            "To document setup considerations such as speed, hardware, privacy, and cost.",
            "To prepare a reusable inference workflow for later labs and projects.",
        ],
        "tools": [
            ("Groq or Replicate", "Hosted inference providers for Llama-family models."),
            ("Ollama or LM Studio", "Local model runners used for on-device deployment."),
            ("Python or CLI client", "Used to send prompts and inspect model responses."),
        ],
        "methodology": [
            "Create an account with a hosted provider such as Groq or Replicate and obtain the required API key.",
            "Select an available Llama model and submit sample prompts using the provider API or playground.",
            "Install Ollama or LM Studio locally and download a compatible Llama model for offline or local inference.",
            "Compare response speed, privacy advantages, setup effort, and hardware dependency between both approaches.",
        ],
        "implementation": [
            "Hosted inference is generally faster to start because it avoids local model download and GPU setup. The user only needs an API key, a chosen model identifier, and a small client snippet to send prompts and read responses.",
            "Local inference, by contrast, gives better privacy and offline availability, but requires model download, runtime installation, and adequate system memory. Ollama simplifies this by exposing a local endpoint, while LM Studio provides a graphical interface for model management and testing.",
            "This lab is conceptually important because it demonstrates deployment flexibility. A student can prototype with hosted inference during early experimentation and later move to local deployment if privacy, latency control, or offline use becomes more important.",
        ],
        "observations": [
            "Hosted Llama access is convenient for quick experimentation and benchmarking.",
            "Local deployment is more suitable when data privacy or offline access matters.",
            "The best platform depends on hardware capacity, turnaround time, and whether the experiment needs a GUI or API-first workflow.",
        ],
        "conclusion": [
            "This lab highlights that open-source LLM usage is not tied to a single platform. Both hosted and local deployment paths are practical, and each offers different operational advantages.",
            "Understanding these deployment options helps students choose the right model-serving approach for future RAG systems, agents, or applications.",
        ],
        "pending_items": [
            "Screenshot or console output from Groq or Replicate showing a successful Llama 3 response.",
            "Screenshot of Ollama or LM Studio local deployment with the selected Llama model.",
            "Any benchmark notes comparing hosted versus local inference.",
        ],
    },
    {
        "lab_no": 9,
        "title": "Experimenting with Open-Source Models on Hugging Face for Zero-Shot Audio Classification, Automatic Speech Recognition, and Text-to-Speech",
        "intro": [
            "This experiment extends the Hugging Face ecosystem workflow from simple NLP tasks toward multimodal audio applications. The manual calls for three task families: zero-shot audio classification, automatic speech recognition, and text-to-speech.",
            "The local repository already contains a Huggingface.ipynb notebook implementing a web-based sentiment analysis app with the Transformers pipeline and Gradio. While the notebook itself focuses on text sentiment rather than audio tasks, it still provides direct evidence that the repository already uses Hugging Face pipelines and a model-backed UI pattern. The audio-specific sections of this report are therefore documented from the manual workflow and Hugging Face pipeline conventions, with execution evidence left pending.",
        ],
        "objectives": [
            "To understand Hugging Face pipeline usage across multiple open-source tasks.",
            "To experiment conceptually with zero-shot audio classification, speech-to-text, and text-to-speech.",
            "To observe how pre-trained open-source models reduce the need for training from scratch.",
            "To document a practical multimodal experimentation workflow.",
        ],
        "tools": [
            ("Hugging Face Transformers", "Provides pre-trained pipelines for inference tasks."),
            ("Gradio", "Supports quick web UI creation for model demos."),
            ("Python notebook environment", "Used to install dependencies and test pipelines."),
            ("Open-source audio models", "Used for audio classification, ASR, and TTS tasks."),
        ],
        "methodology": [
            "Review the existing Hugging Face notebook in the repository to confirm the pipeline-based workflow.",
            "Map the same pipeline concept to the three required audio tasks from the lab manual.",
            "Document the setup steps for loading audio, transcribing speech, and synthesizing speech from text.",
            "Record expected outputs and note what execution proof still needs to be attached later.",
        ],
        "implementation": [
            "The existing notebook shows the core Hugging Face pattern clearly: install transformers and gradio, import the pipeline function, load a task-specific pre-trained model, wrap the inference function inside a Python helper, and expose it through a simple UI. This exact workflow is transferable to other tasks such as zero-shot audio classification, speech recognition, and text-to-speech.",
            "For zero-shot audio classification, the student would typically load an audio classification pipeline, provide an audio sample, and compare the predicted labels with a set of candidate classes. For ASR, the workflow centers on a speech recognition model that converts waveform input into plain text. For TTS, a text input is passed into a synthesis model that generates an audio output file or waveform.",
            "The major educational value of this lab is that the model-loading experience remains consistent even as the modality changes. Once the student understands how to use a pre-trained pipeline and input data correctly, it becomes easier to explore many open-source model categories without training custom systems.",
        ],
        "code_block": [
            "# Existing local pattern from Huggingface.ipynb",
            "from transformers import pipeline",
            "import gradio as gr",
            'classifier = pipeline("sentiment-analysis")',
            "",
            "def analyze_text(text):",
            "    result = classifier(text)[0]",
            '    return f"Label: {result[\'label\']}, Score: {round(result[\'score\'], 4)}"',
        ],
        "observations": [
            "Hugging Face pipelines provide a consistent abstraction layer across tasks, making experimentation quick and approachable.",
            "The main difference between text and audio tasks lies in input preparation and output format, not in the overall inference workflow.",
            "This makes the platform especially useful for labs, where fast prototyping and conceptual understanding are more important than large-scale deployment.",
        ],
        "conclusion": [
            "This lab demonstrates the flexibility of the Hugging Face ecosystem for running diverse open-source model tasks. Even though the repository currently contains direct evidence for a text pipeline rather than audio outputs, the same architectural pattern applies naturally to the required audio experiments.",
            "With the addition of the pending audio outputs, this report can be completed into a full multimodal Hugging Face lab submission.",
        ],
        "pending_items": [
            "Notebook or console output for zero-shot audio classification.",
            "Speech recognition transcript output for an uploaded audio file.",
            "Generated text-to-speech audio sample or screenshot of the synthesis result.",
        ],
    },
    {
        "lab_no": 10,
        "title": "Hands-On Lab Using LangChain and LangFlow",
        "intro": [
            "This experiment introduces orchestration frameworks for building LLM applications. The lab manual sections available locally describe a LangFlow workflow in which a prompt, Hugging Face model component, chain component, and chat I/O blocks are connected visually to form a working conversational pipeline.",
            "The repository does not currently include a saved LangFlow project file, but the manual contains enough detail to document the end-to-end lab procedure and its design logic.",
        ],
        "objectives": [
            "To understand how LangChain concepts appear in a visual LangFlow builder.",
            "To create a simple chatbot pipeline using prompt, model, and chain components.",
            "To extend the workflow with memory for conversational context.",
            "To build a document-questioning flow using file parsing and prompt injection.",
        ],
        "tools": [
            ("LangFlow", "Visual interface for connecting LLM components into flows."),
            ("LangChain concepts", "Prompt templates, chains, memory, and model abstraction."),
            ("Hugging Face model component", "Used in the manual as the selected LLM backend."),
            ("PDF / file component", "Used to support question answering over uploaded documents."),
        ],
        "methodology": [
            "Create a new LangFlow project and rename it appropriately.",
            "Add a Prompt component and define a template with a question variable.",
            "Add a Hugging Face model component and configure the API token and repo ID.",
            "Add a ConversationChain component and connect the prompt and model nodes.",
            "Add Chat Input and Chat Output nodes to allow direct user interaction.",
            "Extend the flow with Chat Memory so that the model can answer follow-up questions.",
            "Build a second flow that loads a PDF, parses data, injects the document into a prompt, and answers questions over the uploaded file.",
        ],
        "implementation": [
            "The manual describes the initial chatbot as a visual chain consisting of Prompt, HuggingFace, and ConversationChain components. Chat Input and Chat Output are then added to make the flow interactive through the LangFlow playground.",
            "The next refinement introduces memory. By including a history variable in the prompt and connecting a Chat Memory component, the flow stores past turns and supports context-aware follow-up questions.",
            "The document-questioning workflow expands the same pattern further. A File component loads the PDF, Parse Data prepares the content, the prompt template injects the document text and the user's question, and the Hugging Face component generates the answer. This effectively demonstrates retrieval-style reasoning at the workflow level even before building a full vector database system.",
        ],
        "observations": [
            "LangFlow is useful for understanding application structure because it makes abstract LLM pipeline components visible and traceable.",
            "Memory and document-input support can be added incrementally without redesigning the whole system.",
            "The visual approach is especially helpful for learning how prompts, models, inputs, and outputs interact in a complete application.",
        ],
        "conclusion": [
            "This lab successfully introduces both the conceptual and practical sides of LangChain-style application building through LangFlow.",
            "The exercise prepares students for later work on RAG systems and agents by showing how LLM application pieces can be composed into a functioning pipeline.",
        ],
        "pending_items": [
            "Screenshot of the LangFlow canvas for the chatbot pipeline.",
            "Screenshot of the memory-enabled flow or playground output.",
            "Screenshot of the document-QA workflow and sample response.",
        ],
    },
    {
        "lab_no": 11,
        "title": "Implementing a RAG-Based LLM Project Using a Vector Database",
        "intro": [
            "This experiment uses retrieval-augmented generation to answer user questions from uploaded documents rather than relying only on the model's internal knowledge. The repository contains two strong local assets for this topic: a HyDE-based RAG notebook inside the RAG folder and a hierarchical rag.py script in the project root.",
            "Together, these files show a practical workflow involving document ingestion, chunking, embedding generation, vector storage, retrieval, context injection, and answer synthesis using a hosted model.",
        ],
        "objectives": [
            "To build a question-answering system grounded in external documents.",
            "To use embeddings and a vector database to retrieve relevant chunks.",
            "To compare standard retrieval with HyDE-style hypothetical document generation.",
            "To understand how RAG improves factual grounding for downstream answers.",
        ],
        "tools": [
            ("Groq API", "Used as the LLM backend in the RAG notebook."),
            ("SentenceTransformers", "Used to generate embeddings for chunks and queries."),
            ("ChromaDB", "Used as the vector database for chunk storage and similarity search."),
            ("PyPDF / text readers", "Used to load source documents for ingestion."),
        ],
        "methodology": [
            "Install the required packages including groq, chromadb, sentence-transformers, and pypdf.",
            "Load documents from user uploads or local files and extract raw text.",
            "Split each document into overlapping chunks suitable for embedding and retrieval.",
            "Encode the chunks with an embedding model and store them in ChromaDB.",
            "Embed incoming user queries, retrieve top-k relevant chunks, and assemble the prompt context.",
            "Generate an answer only from retrieved context and compare this with a HyDE-enhanced retrieval path.",
        ],
        "implementation": [
            "The local notebook shows a complete standard RAG path. It reads uploaded files, chunks the text, encodes the chunks with all-MiniLM-L6-v2 embeddings, and stores them in a Chroma collection. At query time, it embeds the user question, retrieves the most similar chunks, and inserts them into a grounded prompt before calling the Groq-hosted Llama model.",
            "The same notebook also implements HyDE retrieval. Instead of embedding only the original question, it first generates a short hypothetical answer passage, embeds that synthetic text, and uses it to retrieve potentially better matching chunks from the vector database.",
            "This makes the lab especially valuable because it goes beyond a single basic RAG pipeline. It shows how retrieval quality can be improved with query transformation strategies while preserving the core structure of ingestion, embeddings, similarity search, and grounded answer generation.",
        ],
        "code_block": [
            "embedder = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')",
            "collection = chroma_client.create_collection(name='simple_rag')",
            "",
            "def retrieve_standard(query, top_k=4):",
            "    query_embedding = embedder.encode([query]).tolist()[0]",
            "    results = collection.query(query_embeddings=[query_embedding], n_results=top_k)",
            "    return results",
            "",
            "def answer_with_context(user_query, retrieved_docs, model='llama-3.1-8b-instant'):",
            "    ...",
        ],
        "observations": [
            "RAG makes the system more trustworthy because the answer is anchored to retrieved source text rather than pure generation.",
            "The chunking and embedding choices directly affect retrieval quality, so preprocessing is just as important as the final model call.",
            "HyDE can improve retrieval for vague or under-specified questions by generating a semantically richer search representation.",
        ],
        "conclusion": [
            "This lab successfully demonstrates a real RAG workflow using a vector database and an external LLM backend. The repository already contains credible implementation evidence for ingestion, retrieval, and answer generation logic.",
            "The project provides a strong foundation for later document assistants, knowledge-grounded chatbots, and research support systems.",
        ],
        "pending_items": [
            "Any screenshots of the notebook output showing retrieved chunks and final answers.",
            "Optional sample source document used during your actual run.",
        ],
    },
    {
        "lab_no": 12,
        "title": "Building a Simple LLM Agent Using the Phi Data Framework",
        "intro": [
            "This experiment focuses on agent-oriented application design using the Phi Data framework in a browser-based execution environment such as Google Colab. The lab manual clearly defines the expected flow: install the framework, define the agent role and system instructions, run user queries, and observe the resulting behavior.",
            "No Phi Data notebook or script is currently present in the local repository, so the report documents the lab method faithfully and leaves runtime evidence to be attached later.",
        ],
        "objectives": [
            "To understand the structure of a simple LLM agent in Phi Data.",
            "To define an agent role, instructions, and model behavior settings.",
            "To test how an agent responds to different queries.",
            "To evaluate how instruction changes affect output quality and style.",
        ],
        "tools": [
            ("Phi Data framework", "Used to build the agent."),
            ("Google Colab", "Used as the browser-based Python execution environment."),
            ("LLM backend", "Used by the agent to generate responses."),
        ],
        "methodology": [
            "Open Google Colab and create a fresh Python notebook.",
            "Install the Phi Data framework and verify the environment setup.",
            "Define the agent role, task description, and any core behavioral instructions.",
            "Run the agent against multiple user queries and record the responses.",
            "Modify instructions to observe changes in output quality and consistency.",
        ],
        "implementation": [
            "According to the manual, the agent should be implemented as a simple role-driven assistant, such as a research assistant. The student defines instructions, chooses model parameters, and runs test prompts in sequence.",
            "The important educational outcome in this lab is not simply obtaining a response, but observing how an agent differs from a plain prompt-response system. The agent carries explicit behavioral framing and is intended to respond repeatedly under a predefined role and instruction set.",
            "This lab therefore complements the earlier prompting experiments by moving from isolated prompts to a reusable agent abstraction.",
        ],
        "observations": [
            "Agent behavior depends strongly on the clarity and specificity of the system instructions.",
            "Repeated testing with multiple queries helps reveal whether the agent is consistent or brittle.",
            "This lab is especially useful for understanding how prompt engineering evolves into agent design.",
        ],
        "conclusion": [
            "The Phi Data lab introduces the basic structure of an LLM agent and shows how role, instruction, and repeated execution combine to form a more reusable AI workflow.",
            "With execution screenshots attached later, this report can serve as a complete record of the lab procedure and observed behavior.",
        ],
        "pending_items": [
            "Screenshot of Phi Data installation in Colab.",
            "Screenshot of the agent definition code.",
            "Screenshot of agent responses for one or more test queries.",
        ],
    },
    {
        "lab_no": 13,
        "title": "Using LLMs for Code Generation and Bug Detection in Software Development",
        "intro": [
            "This experiment studies how LLMs can support software development in two major ways: generating code from requirements and detecting or correcting bugs in existing code. The lab manual provides concrete programming tasks, prompts, expected outputs, and discussion points for both workflows.",
            "The local repository does not contain a dedicated lab notebook for this topic, but the manual content is detailed enough to produce a faithful lab report with clearly documented expected workflows and outcomes.",
        ],
        "objectives": [
            "To use an LLM for requirement-to-code generation.",
            "To refine prompts for better software structure and input validation.",
            "To detect logical and security bugs with LLM assistance.",
            "To evaluate where LLM help ends and human review remains necessary.",
        ],
        "tools": [
            ("ChatGPT / Gemini / GitHub Copilot", "LLM tool used for code generation and debugging."),
            ("Python", "Recommended language for the lab tasks."),
            ("IDE or notebook environment", "Used to run and verify the generated code."),
        ],
        "methodology": [
            "Provide the student record system problem statement to the LLM and request a complete Python implementation.",
            "Review the generated solution for modularity, assumptions, and missing validations.",
            "Refine the prompt to request OOP design, file storage, and stronger input checks.",
            "Pass the faulty code examples to the LLM and ask it to identify logic and security problems.",
            "Compare the generated fix with manual reasoning before accepting the correction.",
        ],
        "implementation": [
            "The manual's first task uses a student record system example. The LLM is asked to generate code that stores ID, name, and marks, computes the average, and lists students who scored above average. This demonstrates how LLMs can translate structured requirements into a working program skeleton quickly.",
            "The second part focuses on debugging. The manual includes a faulty maximum-finding function that fails for all-negative arrays because the initial maximum is incorrectly set to zero. It also includes a security example involving hardcoded credentials. In both cases, the LLM is prompted to explain the problem and propose corrected code.",
            "This lab is especially important because it shows both the strengths and limits of LLM support. Models can draft code and identify common flaws quickly, but human review is still essential for edge cases, security design, and production readiness.",
        ],
        "code_block": [
            "def find_max(arr):",
            "    max = 0",
            "    for i in arr:",
            "        if i > max:",
            "            max = i",
            "    return max",
            "",
            "Corrected idea:",
            "def find_max(arr):",
            "    max_val = arr[0]",
            "    for i in arr:",
            "        if i > max_val:",
            "            max_val = i",
            "    return max_val",
        ],
        "observations": [
            "LLMs are strong accelerators for draft code generation and first-pass debugging.",
            "Prompt refinement changes the output significantly, especially when requesting OOP structure, validation, or security-conscious design.",
            "Human oversight remains essential because the model may miss hidden requirements or introduce unverified assumptions.",
        ],
        "conclusion": [
            "This lab demonstrates a realistic modern development workflow in which LLMs act as coding assistants, reviewers, and pair programmers.",
            "The most important takeaway is that LLMs improve productivity when combined with careful prompt design and manual verification.",
        ],
        "pending_items": [
            "Screenshots of your actual prompts and model-generated code.",
            "Run output from the corrected code samples, if you want them included later.",
        ],
    },
    {
        "lab_no": 14,
        "title": "Using Multimodal LLMs for Visual Question Answering",
        "intro": [
            "This experiment demonstrates how a multimodal large language model can reason over both image and text inputs. The lab manual provides a complete Gemini-based Visual Question Answering workflow, including SDK installation, API configuration, image upload, image byte conversion, multimodal request construction, and answer logging.",
            "The repository does not yet contain a completed VQA notebook or saved outputs for this lab, so the report uses the manual as the main source and clearly marks the actual response evidence as pending.",
        ],
        "objectives": [
            "To understand multimodal LLMs and the Visual Question Answering task.",
            "To configure the Gemini API for image-plus-text reasoning.",
            "To submit image-based questions and analyze the model responses.",
            "To save the question-answer results for submission and evaluation.",
        ],
        "tools": [
            ("Gemini Multimodal API", "Used to answer questions about the uploaded image."),
            ("google-genai SDK", "Used for Python access to the Gemini model."),
            ("Pillow", "Used to open and convert images into bytes."),
            ("Google Colab", "Used as the execution environment."),
        ],
        "methodology": [
            "Install the required packages such as google-genai and Pillow in Colab.",
            "Generate a Gemini API key and configure the client securely.",
            "Upload an image, inspect it, and convert it to byte format.",
            "Ask one or more visual questions and send the combined text-plus-image request to Gemini.",
            "Save the resulting question-answer pair to CSV for later analysis.",
        ],
        "implementation": [
            "The manual uses a clean multimodal pipeline. After the image is uploaded and loaded through Pillow, it is converted to JPEG bytes using an in-memory buffer. The Gemini client is then initialized, and a request is sent with both the natural-language question and the binary image part.",
            "The sample prompt asks the model to describe the objects present in the image and explain the scene. The workflow can easily be extended with counting questions, color questions, or indoor-versus-outdoor scene interpretation.",
            "This lab is important because it shows that LLM interaction is no longer limited to text. The same application pattern can support educational tools, accessibility assistants, content analysis systems, and visual search interfaces.",
        ],
        "code_block": [
            "from google import genai",
            "from PIL import Image",
            "import io",
            "",
            "client = genai.Client()",
            'question = "Describe the objects present in the image and explain the scene."',
            "response = client.models.generate_content(",
            '    model="gemini-2.5-flash",',
            "    contents=[question, genai.types.Part.from_bytes(data=image_bytes, mime_type='image/jpeg')]",
            ")",
            "print(response.text)",
        ],
        "observations": [
            "The lab highlights how multimodal models combine visual context with linguistic reasoning.",
            "The quality of the answer depends on both image clarity and question specificity.",
            "Structured logging through CSV makes it easier to compare different questions and different images systematically.",
        ],
        "conclusion": [
            "This experiment demonstrates a practical VQA workflow using Gemini, Python, and Colab. It provides a clear example of how modern multimodal LLMs extend far beyond text-only generation.",
            "Once the actual image outputs and responses are attached, the report becomes a complete record of the experiment.",
        ],
        "pending_items": [
            "Uploaded image screenshot or sample image used in the VQA run.",
            "Notebook output showing Gemini's response to the visual question.",
            "Saved CSV or screenshot proving the output logging step.",
        ],
    },
    {
        "lab_no": 15,
        "title": "Course Project Using Llama or Gemini with a Streamlit or Gradio Frontend",
        "intro": [
            "This final experiment asks for a compact end-to-end LLM application with a user-facing frontend. The strongest reusable local asset for this requirement is Huggingface.ipynb, which already implements a Gradio-based sentiment analysis application using a Hugging Face transformer pipeline.",
            "Although the manual mentions Llama or Gemini with Streamlit or Gradio, the local project evidence most closely matches the Gradio frontend path. This report therefore documents the existing Gradio app as the core course-project style artifact and notes where future extensions toward Gemini or Llama would fit.",
        ],
        "objectives": [
            "To build a simple user-facing LLM application with a frontend interface.",
            "To connect a model pipeline to an interactive web UI.",
            "To demonstrate end-to-end flow from user input to model output.",
            "To document how the project can be extended into a larger course deliverable.",
        ],
        "tools": [
            ("Hugging Face Transformers", "Used to load the sentiment-analysis pipeline."),
            ("Gradio", "Used to create the frontend interface."),
            ("Python notebook", "Used to run the application setup and launch sequence."),
        ],
        "methodology": [
            "Install the required packages transformers and gradio.",
            "Load a pre-trained model pipeline for inference.",
            "Wrap the model call in a helper function that formats the result for the UI.",
            "Create a Gradio interface with text input and text output components.",
            "Launch the interface and validate it with sample inputs.",
        ],
        "implementation": [
            "The local Huggingface.ipynb notebook already implements a compact AI web application. It loads a sentiment-analysis pipeline, defines an analyze_text function, and connects that function to a Gradio interface titled 'My First AI Sentiment App'. The user enters a sentence, and the application returns the predicted label and score.",
            "This is a credible frontend project because it demonstrates the complete user journey: input capture, backend model inference, formatted output rendering, and browser-based interaction through Gradio. The same structure could later be adapted to Gemini or Llama by replacing the backend inference call while preserving the UI layer.",
            "As a course project, this application can be extended with history, multiple tasks, richer UI components, or deployment. Even in its simple form, it satisfies the key educational requirement of combining a model with a real frontend.",
        ],
        "code_block": [
            "from transformers import pipeline",
            "import gradio as gr",
            "",
            'classifier = pipeline("sentiment-analysis")',
            "",
            "def analyze_text(text):",
            "    result = classifier(text)[0]",
            '    return f"Label: {result[\'label\']}, Score: {round(result[\'score\'], 4)}"',
            "",
            "demo = gr.Interface(fn=analyze_text, inputs='text', outputs='text',",
            "                    title='My First AI Sentiment App',",
            "                    description='Type a sentence below to see if it is Positive or Negative!')",
            "demo.launch(share=True)",
        ],
        "observations": [
            "A small Gradio application is enough to demonstrate the essential architecture of an LLM-powered project.",
            "Frontend wrappers increase the accessibility of model experiments because they remove the need for direct code interaction by end users.",
            "This project can scale naturally into a larger semester deliverable by replacing the backend task or expanding the interface.",
        ],
        "conclusion": [
            "This lab successfully maps the course-project requirement to an existing repo-backed Gradio application. The project demonstrates that even a compact model demo can serve as a valid end-to-end AI application when it includes both inference and a usable frontend.",
            "The same structure can later be upgraded to a Gemini or Llama backend if a broader course showcase is needed.",
        ],
        "pending_items": [
            "Screenshot of the Gradio interface while running.",
            "Sample input and output screenshot from the live app.",
            "Any deployment link or enhanced project version you want included later.",
        ],
    },
]


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = []
    for report in REPORTS:
        outputs.append(write_report(**report))
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
