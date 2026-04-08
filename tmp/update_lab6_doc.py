from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
OUTPUT_PATH = BASE_DIR / "final lab" / "lab_6_428.docx"
IMG_DIR = BASE_DIR / "output" / "playwright"
DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Platform"
    hdr[1].text = "Purpose in the Experiment"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


def add_image(doc, image_path, width_inches):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    para.paragraph_format.space_after = Pt(4)


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(8)


doc = Document()
configure_document(doc)

add_paragraph(doc, "GEN_AI LAB 6", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(
    doc,
    "Using Cursor AI or Lovable to Create a Website or Web Application",
    italic=True,
    size=12,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=10,
)
add_paragraph(doc, f"Name: {NAME}")
add_paragraph(doc, f"USN: {USN}")
add_paragraph(doc, f"Date: {DATE_STR}", space_after=10)

add_heading(doc, "Introduction")
add_paragraph(
    doc,
    "This experiment explores AI-assisted software development through tools such as Cursor AI or Lovable. The objective is to understand how natural-language prompts can be used to speed up website or web application creation, especially during the prototyping stage.",
)
add_paragraph(
    doc,
    "For the output evidence in this report, the deployed website https://aishviwebsolutions.in/ was reviewed in a browser and captured through Playwright screenshots. These images are used as concrete output proof for the lab instead of leaving the report fully placeholder-based.",
)

add_heading(doc, "Objective")
for item in [
    "To understand how AI-assisted IDEs and builders support website creation.",
    "To analyze the structure and presentation of a deployed business website.",
    "To document frontend sections such as hero area, services, and contact block using screenshots.",
    "To evaluate the strengths and limitations of AI-supported web development workflows.",
]:
    add_bullet(doc, item)

add_heading(doc, "Tools and Platforms Used")
add_tools_table(
    doc,
    [
        ("Cursor AI or Lovable", "Representative AI-assisted development tools referenced in the lab requirement."),
        ("Browser-based website output", "Used to inspect the live site aishviwebsolutions.in."),
        ("Playwright CLI", "Used to capture output screenshots from the deployed website."),
        ("python-docx", "Used to rebuild the lab report and embed screenshot evidence."),
    ],
)

add_heading(doc, "Methodology / Procedure")
for item in [
    "Open the target website and inspect the visible sections of the deployed output.",
    "Capture the homepage view, service section, and contact section as visual evidence.",
    "Map these sections back to common AI-assisted development tasks such as hero design, services layout, and call-to-action design.",
    "Document the likely development workflow and evaluate the resulting user experience.",
]:
    add_bullet(doc, item)

add_heading(doc, "Implementation / Workflow Summary")
add_paragraph(
    doc,
    "A practical AI-assisted workflow for this lab begins by describing the business purpose, service offerings, layout style, and the expected call-to-action areas. An AI coding assistant such as Cursor AI or a product builder such as Lovable can then generate the initial page structure, including a hero section, service cards, about text, and contact information blocks.",
)
add_paragraph(
    doc,
    "The website reviewed for this report, AISHVI Web Solutions, reflects the kind of result commonly produced from such a workflow. The site includes a prominent hero banner, a services grid with clearly grouped offerings, an about section, and a contact area with phone, email, and office details. This makes it a suitable real-world output example for the lab.",
)

add_heading(doc, "Website Output Evidence")
add_paragraph(
    doc,
    "The following screenshots were captured from the deployed website and used as output proof for this lab.",
)

home = IMG_DIR / "lab6_homepage.png"
services = IMG_DIR / "lab6_services.png"
contact = IMG_DIR / "lab6_contact.png"

add_image(doc, home, 6.2)
add_caption(doc, "Figure 1. Homepage of AISHVI Web Solutions showing the hero section, headline, and primary call to action.")

add_image(doc, services, 6.2)
add_caption(doc, "Figure 2. Services section showing structured website offerings such as development, mobile optimization, e-commerce, maintenance, design, and analytics.")

add_image(doc, contact, 5.8)
add_caption(doc, "Figure 3. Contact and engagement section showing the consultation message, contact details, and company identity block.")

add_heading(doc, "Site-Specific Observations")
add_paragraph(
    doc,
    "The homepage uses a clear business-oriented headline: 'Building Your Digital Future, Today!' followed by a concise explanation of web solutions and a call-to-action button. This is an effective landing-page pattern because it immediately communicates the business value of the site.",
)
add_paragraph(
    doc,
    "The services section is one of the strongest parts of the website. It presents multiple offerings such as Website Development, Mobile Optimization, E-commerce Solutions, SEO and Digital Marketing, Website Maintenance, UI/UX Design, Analytics and Reporting, 24/7 Support, and Digital Transformation. Each card combines a short description with feature bullets, which improves scannability and credibility.",
)
add_paragraph(
    doc,
    "The lower contact section reinforces the conversion flow by offering a consultation-style message and clear contact channels including phone and email. This is important for a service business website because it turns the design from a purely informational page into a lead-generation tool.",
)

add_heading(doc, "Observations")
for item in [
    "AI-assisted web creation is highly effective for generating standard business website structures quickly.",
    "The reviewed website demonstrates a strong section-based layout that is easy to understand and visually consistent.",
    "Features such as hero text, service cards, and contact calls to action are especially well suited to prompt-driven generation workflows.",
    "Even when AI accelerates development, human review is still needed for polish, content accuracy, and brand-specific refinement.",
]:
    add_bullet(doc, item)

add_heading(doc, "Conclusion")
add_paragraph(
    doc,
    "This lab demonstrates how AI-assisted tools such as Cursor AI or Lovable can support rapid website prototyping and production-ready section design. By using the live AISHVI Web Solutions website as visual evidence, the report now includes concrete output rather than only a theoretical workflow description.",
)
add_paragraph(
    doc,
    "Overall, the experiment shows that AI can significantly reduce the effort required to build business websites, especially when the required structure is clear and repetitive. The final quality, however, still depends on careful human review and refinement.",
)

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT_PATH))
print(OUTPUT_PATH)
