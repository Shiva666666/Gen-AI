from __future__ import annotations

import json
import re
from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"

IMAGE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


@dataclass(frozen=True)
class SectionSpec:
    number: int
    heading_key: str
    final_heading: str
    source_path: Path


ROOT = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
TEMPLATE_PATH = Path(r"C:\Users\licha\Downloads\Lab Manual_CS3240.docx")
SOURCE_DIR = ROOT / "final lab"
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "GenAI_Final_Lab_Manual_1RVU23CSE428.docx"
SUMMARY_PATH = OUTPUT_DIR / "GenAI_Final_Lab_Manual_1RVU23CSE428.summary.json"

STUDENT_NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"
SEMESTER = "6th Semester"


SECTION_SPECS = [
    SectionSpec(1, "exp1", "Exp 1 : Foundational Interaction & Multimodal Generative AI", SOURCE_DIR / "lab_1_428.docx"),
    SectionSpec(2, "lab2", "Lab 2 – Using Kaggle, Fast AI, Hugging Face", SOURCE_DIR / "lab_2_428.docx"),
    SectionSpec(3, "lab3", "Lab 3 – Implementation and Visualization of Word2Vec Embeddings", SOURCE_DIR / "lab_3_428.docx"),
    SectionSpec(4, "lab4", "Lab 4 Logical Reasoning & Structural Prompting", SOURCE_DIR / "lab_4_428.docx"),
    SectionSpec(5, "lab5", "Lab 5. Using ChatGPT/Gemini: Generate your resume, Simulate a complete interview", SOURCE_DIR / "lab_5_428.docx"),
    SectionSpec(6, "lab6", "Lab 6. Using Cursor AI/Lovable: Create a website, web app", SOURCE_DIR / "lab_6_428.docx"),
    SectionSpec(7, "lab7", "Lab 7 – Gemini-Pro: Generative API Key and Accessing the Model using API", SOURCE_DIR / "lab_7_428.docx"),
    SectionSpec(8, "lab8", "Lab 8 – Using Meta’s LLaMA-3 Models (Cloud & Local Deployment)", SOURCE_DIR / "lab_8_428.docx"),
    SectionSpec(9, "lab9", "Lab 9. Experimenting open-source models with hugging face (zero shot audio classification, automatic speech recognition, Text to speech)", SOURCE_DIR / "lab_9_428.docx"),
    SectionSpec(10, "lab10", "Lab 10.: Demonstrate an experiment with LangChain and LangFlow", SOURCE_DIR / "lab_10_428.docx"),
    SectionSpec(11, "lab11", "Lab 11: Experiment with RAG and VectorDb", SOURCE_DIR / "lab_11_428.docx"),
    SectionSpec(12, "lab12", "Lab 12. Build a Simple LLM Agent using phi data framework", SOURCE_DIR / "lab_12_428.docx"),
    SectionSpec(13, "lab13", "Lab 13 – Use LLMs for Code Generation and Bug Detection in Software Development", SOURCE_DIR / "lab_13_428.docx"),
    SectionSpec(14, "lab14", "Lab 14. – Using Multimodal Large Language Models (Gemini) for Visual Question Answering", SOURCE_DIR / "lab_14_428.docx"),
    SectionSpec(15, "lab15", "Lab 15 – Course Project Using a Free Local VQA Model with a Gradio Frontend", SOURCE_DIR / "lab_15_428.docx"),
]


SECTION_HEADING_PATTERNS = {
    "exp1": re.compile(r"^Exp\s*1\b", re.I),
    "lab2": re.compile(r"^Lab\s*2\b", re.I),
    "lab3": re.compile(r"^Lab\s*3\b", re.I),
    "lab4": re.compile(r"^Lab\s*4\b", re.I),
    "lab5": re.compile(r"^Lab\s*5\b", re.I),
    "lab6": re.compile(r"^Lab\s*6\b", re.I),
    "lab7": re.compile(r"^Lab\s*7\b", re.I),
    "lab8": re.compile(r"^Lab\s*8\b", re.I),
    "lab9": re.compile(r"^Lab\s*9\b", re.I),
    "lab10": re.compile(r"^Lab\s*10\b", re.I),
    "lab11": re.compile(r"^Lab\s*11\b", re.I),
    "lab12": re.compile(r"^Lab\s*12\b", re.I),
    "lab13": re.compile(r"^Lab\s*13\b", re.I),
    "lab14": re.compile(r"^Lab\s*14\b", re.I),
}

SUBSECTION_STYLE_MAP = {
    "introduction": "Heading 2",
    "overview": "Heading 2",
    "objective": "Heading 2",
    "aim": "Heading 2",
    "procedure": "Heading 2",
    "methodology": "Heading 2",
    "methodology / procedure": "Heading 2",
    "implementation summary": "Heading 2",
    "implementation / workflow summary": "Heading 2",
    "working procedure": "Heading 2",
    "tools and frameworks": "Heading 2",
    "tools and platforms used": "Heading 2",
    "tools and libraries used": "Heading 2",
    "results and analysis": "Heading 2",
    "observation": "Heading 2",
    "observations": "Heading 2",
    "conclusion": "Heading 2",
    "result": "Heading 2",
    "result / conclusion": "Heading 2",
    "generated interview simulation": "Heading 2",
    "representative notebook snippet": "Heading 2",
    "representative snippets": "Heading 2",
    "screenshot evidence": "Heading 2",
    "notebook output evidence": "Heading 2",
    "website output evidence": "Heading 2",
    "output evidence from lm studio screenshots": "Heading 2",
    "dataset preparation": "Heading 2",
    "training logic": "Heading 2",
    "workflow highlights": "Heading 2",
    "concept": "Heading 2",
    "implementation / workflow summary": "Heading 2",
}

PART_HEADING_RE = re.compile(r"^Part\s+[A-D]:", re.I)
METADATA_RE = re.compile(r"^(Name|USN|Date)\s*:", re.I)
SKIP_PARAGRAPH_PATTERNS = [
    re.compile(r"^Notebook Output Evidence$", re.I),
    re.compile(r"^Website Output Evidence$", re.I),
    re.compile(r"^Screenshot Evidence$", re.I),
    re.compile(r"^Output Evidence from LM Studio Screenshots$", re.I),
    re.compile(r"\bevidence base for this report\b", re.I),
    re.compile(r"\bthe report therefore\b", re.I),
    re.compile(r"\bthe report now\b", re.I),
    re.compile(r"\boutput proof\b", re.I),
    re.compile(r"\boutput evidence\b", re.I),
    re.compile(r"\bevidence source\b", re.I),
    re.compile(r"\bprimary basis\b", re.I),
    re.compile(r"^The following screenshots\b", re.I),
    re.compile(r"^The uploaded notebook does not include\b", re.I),
    re.compile(r"^Lab 11 has now been fully reoriented\b", re.I),
]
REWRITE_RULES = [
    (
        re.compile(r"^Overview:\s*This report documents.+$", re.I),
        "Overview: This lab covers three practical AI workflows that move from a ready-to-use transformer pipeline to transfer learning in computer vision and data-driven salary analysis.",
    ),
    (
        re.compile(r"^For this report, the workflow is documented.+$", re.I),
        "The workflow includes a complete sample interview simulation that matches the stated lab objective and shows how an LLM can be used to rehearse both technical and HR-style questions.",
    ),
    (
        re.compile(r"^For the output evidence in this report,\s*", re.I),
        "",
    ),
    (
        re.compile(r"^The website reviewed for this report,\s*", re.I),
        "",
    ),
    (
        re.compile(r"^The supplied screenshot\b", re.I),
        "The screenshot",
    ),
    (
        re.compile(r"^The screenshots supplied for this lab\b", re.I),
        "The screenshots",
    ),
    (
        re.compile(r"^This experiment focuses on building hands-on LLM workflows using LangChain concepts such as prompts, chains, output parsing, and tool-calling agents\..*$", re.I),
        "This experiment focuses on building hands-on LLM workflows using LangChain concepts such as prompts, chains, output parsing, and tool-calling agents.",
    ),
    (
        re.compile(r"^The uploaded notebook begins\b", re.I),
        "The workflow begins",
    ),
    (
        re.compile(r"^This report is rebuilt around the Kaggle notebook 'Starter LLM RAG implementation' by Vinayak Tiwari\..*$", re.I),
        "This lab follows the Kaggle notebook 'Starter LLM RAG implementation' by Vinayak Tiwari and uses its RAG pipeline, explanations, and visible outputs.",
    ),
    (
        re.compile(r"^This lab successfully demonstrates a working local LLM deployment and inference workflow using LM Studio\..*$", re.I),
        "This lab demonstrates a working local LLM deployment and inference workflow using LM Studio.",
    ),
    (
        re.compile(r"^This lab report now uses actual outputs from the uploaded notebook instead of placeholder text\.\s*", re.I),
        "",
    ),
    (
        re.compile(r"^The lab was completed successfully using a free local coding model\..*$", re.I),
        "The lab was completed using a free local coding model with outputs for code creation, bug analysis, and corrected-code verification.",
    ),
    (
        re.compile(r"^The lab was completed successfully using a free Hugging Face VQA model\..*$", re.I),
        "The lab was completed using a free Hugging Face VQA model with the actual image and answer summary from the executed workflow.",
    ),
    (
        re.compile(r"^Lab 15 was successfully implemented as a free local Gradio-based VQA project\..*$", re.I),
        "Lab 15 was successfully implemented as a free local Gradio-based VQA project with a visible application screenshot and predictions.",
    ),
]


def iter_block_items(parent: DocxDocument | _Cell) -> Iterable[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def append_body_element(body_elm, new_elm) -> None:
    if len(body_elm) and body_elm[-1].tag == qn("w:sectPr"):
        body_elm.insert(len(body_elm) - 1, new_elm)
    else:
        body_elm.append(new_elm)


def clone_paragraph_prototype(paragraph: Paragraph):
    return deepcopy(paragraph._element)


def paragraph_text(block: Paragraph | Table) -> str:
    if isinstance(block, Paragraph):
        return block.text.strip()
    return " ".join(cell.text.strip() for row in block.rows for cell in row.cells).strip()


def paragraph_has_drawings(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//*[local-name()='drawing']"))


def is_metadata_table(table: Table) -> bool:
    keys = []
    for row in table.rows[:4]:
        if not row.cells:
            continue
        keys.append(row.cells[0].text.strip().lower())
    return any(key in {"name", "usn", "date", "title"} for key in keys)


def clone_block_into_doc(block: Paragraph | Table, dest_doc: DocxDocument):
    new_elm = deepcopy(block._element)
    remap_relationships(new_elm, block.part, dest_doc.part)
    refresh_drawing_ids(new_elm, dest_doc.part)
    append_body_element(dest_doc._body._element, new_elm)
    return new_elm


def remap_relationships(element, source_part, dest_part) -> None:
    rid_attrs = [qn("r:embed"), qn("r:id"), qn("r:link")]

    for node in element.iter():
        for attr in rid_attrs:
            old_rid = node.attrib.get(attr)
            if not old_rid:
                continue
            rel = source_part.rels.get(old_rid)
            if rel is None:
                continue
            if rel.reltype == IMAGE_REL:
                new_rid, _ = dest_part.get_or_add_image(BytesIO(rel.target_part.blob))
            elif rel.reltype == HYPERLINK_REL:
                new_rid = dest_part.relate_to(rel.target_ref, rel.reltype, is_external=True)
            elif rel.is_external:
                new_rid = dest_part.relate_to(rel.target_ref, rel.reltype, is_external=True)
            else:
                new_rid = dest_part.relate_to(rel.target_part, rel.reltype)
            node.set(attr, new_rid)


def refresh_drawing_ids(element, dest_part) -> None:
    for node in element.iter():
        if node.tag == f"{{{WP_NS}}}docPr":
            node.set("id", str(dest_part.next_id))
        elif node.tag == f"{{{PIC_NS}}}cNvPr":
            node.set("id", str(dest_part.next_id))


def normalize_heading(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("–", "-").replace("—", "-")).strip().lower()


def section_heading_paragraphs(template_doc: DocxDocument) -> dict[str, Paragraph]:
    found: dict[str, Paragraph] = {}
    for paragraph in template_doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        for key, pattern in SECTION_HEADING_PATTERNS.items():
            if key not in found and pattern.match(text):
                found[key] = paragraph
    return found


def fill_front_matter(template_doc: DocxDocument) -> None:
    for paragraph in template_doc.paragraphs:
        text = paragraph.text.strip()
        if "This is to certify that Mr./Ms." in text:
            paragraph.text = "This is to certify that Mr./Ms. Shiva Dhanush S has"
        elif "USN:" in text and "Semester:" in text and "--------------------------------" in text:
            paragraph.text = f"USN: {USN}    Semester: {SEMESTER}"


def strip_template_labs(template_doc: DocxDocument, first_heading_paragraph: Paragraph) -> None:
    body_elm = template_doc._body._element
    first_elm = first_heading_paragraph._element
    removing = False
    for child in list(body_elm):
        if child is first_elm:
            removing = True
        if removing and child.tag != qn("w:sectPr"):
            body_elm.remove(child)


def add_page_break(doc: DocxDocument) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def dedupe_adjacent_blank_paragraphs(doc: DocxDocument) -> None:
    body_elm = doc._body._element
    previous_blank = False
    for child in list(body_elm):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            is_blank = not paragraph.text.strip() and not paragraph_has_drawings(paragraph)
            if is_blank and previous_blank:
                body_elm.remove(child)
                continue
            previous_blank = is_blank
        else:
            previous_blank = False


def should_skip_block(block: Paragraph | Table, section: SectionSpec, started: bool, state: dict) -> bool:
    if isinstance(block, Table):
        return is_metadata_table(block)

    text = block.text.strip()
    has_image = paragraph_has_drawings(block)

    if not started:
        if section.number == 2:
            if text or has_image:
                state["seen_any"] = state.get("seen_any", 0) + 1
            if state.get("seen_any") == 1 and (text or has_image):
                state["copy_after_title"] = True
                return True
            return False

        if METADATA_RE.match(text):
            if text.lower().startswith("date:"):
                state["copy_after_metadata"] = True
            return True
        return True

    if not text and not has_image:
        return True
    if METADATA_RE.match(text):
        return True
    if section.number == 2 and state.get("drop_next_blank_after_title") and not text and not has_image:
        return True
    if PART_HEADING_RE.match(text):
        return False
    return False


def apply_subsection_style(paragraph: Paragraph) -> None:
    text = paragraph.text.strip().rstrip(":").lower()
    if not text:
        return
    if PART_HEADING_RE.match(paragraph.text.strip()):
        paragraph.style = "Heading 3"
        return
    if text in SUBSECTION_STYLE_MAP:
        paragraph.style = SUBSECTION_STYLE_MAP[text]


def cleanup_paragraph(paragraph: Paragraph) -> bool:
    original_text = paragraph.text
    text = original_text.strip()
    if not text:
        return True

    for pattern, replacement in REWRITE_RULES:
        if pattern.search(text):
            text = pattern.sub(replacement, text).strip()

    if not text:
        paragraph._element.getparent().remove(paragraph._element)
        return False

    for pattern in SKIP_PARAGRAPH_PATTERNS:
        if pattern.search(text):
            paragraph._element.getparent().remove(paragraph._element)
            return False

    if text != original_text.strip():
        paragraph.text = text
    return True


def copy_section_content(dest_doc: DocxDocument, section: SectionSpec) -> dict:
    source_doc = Document(section.source_path)
    blocks_added = 0
    images_added = 0
    state: dict[str, object] = {}
    started = False

    for block in iter_block_items(source_doc):
        if not started and (state.get("copy_after_metadata") or state.get("copy_after_title")):
            started = True
            state.pop("copy_after_metadata", None)
            state.pop("copy_after_title", None)

        if should_skip_block(block, section, started, state):
            continue

        if not started:
            continue

        if isinstance(block, Paragraph):
            image_count = len(block._element.xpath(".//*[local-name()='blip']"))
            images_added += image_count
        new_elm = clone_block_into_doc(block, dest_doc)
        blocks_added += 1

        if isinstance(block, Paragraph):
            new_paragraph = Paragraph(new_elm, dest_doc._body)
            if not cleanup_paragraph(new_paragraph):
                continue
            apply_subsection_style(new_paragraph)

    return {"blocks_added": blocks_added, "displayed_images_added": images_added}


def append_heading(doc: DocxDocument, prototype_elm, text: str):
    heading_elm = deepcopy(prototype_elm)
    append_body_element(doc._body._element, heading_elm)
    paragraph = Paragraph(heading_elm, doc._body)
    paragraph.text = text
    return paragraph


def count_displayed_images(doc: DocxDocument) -> int:
    return len(doc._body._element.xpath(".//*[local-name()='blip']"))


def first_heading_key_and_paragraph(heading_map: dict[str, Paragraph]) -> tuple[str, Paragraph]:
    order = list(SECTION_HEADING_PATTERNS.keys())
    for key in order:
        if key in heading_map:
            return key, heading_map[key]
    raise RuntimeError("No section headings found in template.")


def build_manual() -> dict:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE_PATH)

    heading_map = section_heading_paragraphs(doc)
    missing = [key for key in SECTION_HEADING_PATTERNS if key not in heading_map]
    if missing:
        raise RuntimeError(f"Template headings not found: {missing}")

    heading_prototypes = {key: clone_paragraph_prototype(par) for key, par in heading_map.items()}
    fill_front_matter(doc)

    _, first_heading = first_heading_key_and_paragraph(heading_map)
    strip_template_labs(doc, first_heading)

    section_summaries = []
    for index, section in enumerate(SECTION_SPECS):
        add_page_break(doc)
        if section.number == 15:
            prototype = heading_prototypes["lab14"]
        else:
            prototype = heading_prototypes[section.heading_key]
        append_heading(doc, prototype, section.final_heading)
        summary = copy_section_content(doc, section)
        summary["section"] = section.final_heading
        summary["source"] = str(section.source_path)
        section_summaries.append(summary)

    dedupe_adjacent_blank_paragraphs(doc)
    doc.save(OUTPUT_PATH)

    final_doc = Document(OUTPUT_PATH)
    summary = {
        "output_path": str(OUTPUT_PATH),
        "student_name": STUDENT_NAME,
        "usn": USN,
        "semester": SEMESTER,
        "section_count": len(SECTION_SPECS),
        "displayed_image_count": count_displayed_images(final_doc),
        "sections": section_summaries,
    }
    SUMMARY_PATH.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


if __name__ == "__main__":
    result = build_manual()
    print(json.dumps(result, indent=2))
