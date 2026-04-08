from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
OUTPUT_DIR = BASE_DIR / "final lab"
ASSET_DIR = BASE_DIR / "lab3_assets"
OUTPUT_PATH = OUTPUT_DIR / "lab_3_428.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)
    return para


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return para


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(8)


def add_image(doc, image_path, width_inches):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

normal_style = doc.styles["Normal"]
normal_style.font.name = "Times New Roman"
normal_style.font.size = Pt(12)

add_paragraph(doc, "LAB-3 Report: Implementation and Visualization of Word2Vec Embeddings", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_paragraph(doc, "Name: Shiva Dhanush S", size=12)
add_paragraph(doc, "USN: 1RVU23CSE428", size=12)
add_paragraph(doc, "Date: 08-04-2026", size=12, space_after=10)

add_heading(doc, "Introduction")
add_paragraph(
    doc,
    "This laboratory exercise focuses on building semantic word embeddings with the Word2Vec algorithm and visualizing the learned relationships between words. Word embeddings transform text into dense numerical vectors, allowing a machine learning model to capture context, similarity, and usage patterns in a compact mathematical form. Instead of treating each word as an isolated symbol, Word2Vec maps related terms close to one another in vector space, which makes it a foundational method in natural language processing.",
)
add_paragraph(
    doc,
    "In this experiment, a small NLP-oriented corpus is prepared, tokenized, and used to train a Word2Vec model through the Gensim library. The generated 100-dimensional vectors are then reduced to two dimensions with Principal Component Analysis so that the embeddings can be interpreted visually. The final plot acts as an intuitive summary of how the model organizes word meaning from the sample text.",
)

add_heading(doc, "Objective")
add_bullet(doc, "To understand how Word2Vec learns distributed word representations from a text corpus.")
add_bullet(doc, "To preprocess text data using tokenization and normalization techniques.")
add_bullet(doc, "To train a Word2Vec model and inspect the generated embedding vectors.")
add_bullet(doc, "To reduce high-dimensional vectors to two dimensions using PCA and visualize them clearly.")

add_heading(doc, "Tools and Libraries Used")
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Library / Tool"
hdr[1].text = "Purpose in the Experiment"
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    set_cell_shading(cell, "D9EAF7")

rows = [
    ("Python", "Main programming language used for implementation."),
    ("Gensim", "Training the Word2Vec model and retrieving word vectors."),
    ("NLTK", "Sentence tokenization and preprocessing of the text corpus."),
    ("Scikit-learn", "Applying PCA for dimensionality reduction."),
    ("Matplotlib", "Plotting the reduced embeddings in a two-dimensional graph."),
]
for left, right in rows:
    cells = table.add_row().cells
    cells[0].text = left
    cells[1].text = right
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

doc.add_paragraph()

add_heading(doc, "Methodology")
add_bullet(doc, "Corpus preparation: A small set of sentences related to natural language processing is defined as the input dataset.")
add_bullet(doc, "Text normalization: Each sentence is converted to lowercase so that differently cased words are treated consistently.")
add_bullet(doc, "Tokenization: NLTK splits the sentences into individual words, creating the tokenized corpus required by Word2Vec.")
add_bullet(doc, "Model training: Word2Vec is trained with a vector size of 100 and a suitable context window to learn semantic relationships.")
add_bullet(doc, "Vector extraction: A subset of vocabulary items is selected and their embedding vectors are printed for inspection.")
add_bullet(doc, "Dimensionality reduction: PCA reduces the 100-dimensional vectors into 2 components for visual interpretation.")
add_bullet(doc, "Visualization: The reduced vectors are plotted and annotated to show the relative positions of selected words.")

add_heading(doc, "Implementation Summary")
add_paragraph(
    doc,
    "The first stage of implementation installs the required libraries and prepares the sample corpus. NLTK resources such as tokenizers are downloaded, after which the corpus is converted to lowercase and transformed into token lists. This tokenized representation becomes the input to the Word2Vec model. Once training is complete, the model is saved and reloaded so that the embeddings can be accessed for further analysis.",
)
add_paragraph(
    doc,
    "The second stage applies PCA to the selected word vectors. Since 100-dimensional vectors cannot be interpreted directly on paper, PCA projects them into a two-dimensional coordinate system while retaining as much of the original structure as possible. The resulting coordinates are plotted with Matplotlib, and each point is labeled with its corresponding token to make the semantic relationships easier to understand.",
)

add_image(doc, ASSET_DIR / "page2_img1.png", 6.2)
add_caption(doc, "Figure 1. Word2Vec model creation, tokenization, and vector extraction.")
add_image(doc, ASSET_DIR / "page2_img2.png", 5.2)
add_caption(doc, "Figure 2. PCA-based dimensionality reduction and 2D plotting code.")

add_heading(doc, "Results and Analysis")
add_paragraph(
    doc,
    "The output confirms that the trained model generates numerical embeddings for words in the vocabulary. Each token is represented by a floating-point vector whose values encode the contextual patterns learned from the input sentences. Although the corpus in this experiment is small, the vectors still demonstrate how words that appear in similar contexts can acquire meaningful geometric relationships.",
)
add_paragraph(
    doc,
    "The PCA visualization provides a simplified map of the embedding space. Words such as 'word', 'tool', and 'creating' appear positioned near conceptually relevant neighbors, while other terms are separated according to how they occur in the corpus. This visual arrangement helps verify that the model is not simply memorizing tokens but is instead learning a structured representation of the sample language data.",
)
add_paragraph(
    doc,
    "The experiment also highlights an important practical insight: the quality of embeddings depends strongly on the size and richness of the corpus. With a larger dataset, Word2Vec would capture deeper semantic and syntactic patterns. Even so, this small-scale implementation is effective for understanding the complete workflow of preprocessing text, training embeddings, extracting vectors, and visualizing relationships.",
)

add_image(doc, ASSET_DIR / "page3_img1.png", 6.5)
add_caption(doc, "Figure 3. Console output showing tokenized corpus and generated embedding vectors.")
add_image(doc, ASSET_DIR / "page3_img2.png", 6.0)
add_caption(doc, "Figure 4. Two-dimensional visualization of selected word embeddings after PCA.")

add_heading(doc, "Conclusion")
add_paragraph(
    doc,
    "This lab successfully demonstrated the implementation of Word2Vec embeddings using a simple NLP corpus. The workflow covered preprocessing, training, saving and loading the model, inspecting the learned vectors, and visualizing them after dimensionality reduction. The results show how Word2Vec converts words into meaningful numerical representations that can support downstream language processing tasks.",
)
add_paragraph(
    doc,
    "Overall, the experiment provides a strong foundation for more advanced language models and embedding techniques. By understanding this pipeline, it becomes easier to explore larger corpora, experiment with different model settings, and apply embeddings in applications such as document classification, similarity search, clustering, and conversational AI systems.",
)

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT_PATH))
print(OUTPUT_PATH)
