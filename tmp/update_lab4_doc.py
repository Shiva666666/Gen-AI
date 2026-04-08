from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\licha\OneDrive\Desktop\Gen-AI")
OUTPUT_PATH = BASE_DIR / "final lab" / "lab_4_428.docx"
DATE_STR = "08-04-2026"
NAME = "Shiva Dhanush S"
USN = "1RVU23CSE428"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    para.paragraph_format.space_after = Pt(3)


def add_code_line(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10.5)


def add_tools_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Tool / Source"
    hdr[1].text = "Use in Lab 4"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        set_cell_shading(cell, "D9EAF7")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_paragraph()


doc = Document()
configure_document(doc)

add_paragraph(doc, "GEN_AI LAB 4", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(
    doc,
    "Prompt Engineering Practice: Chain of Thought, Tabular, Fill-in-the-Blank, RGC, Zero-Shot, One-Shot, and Few-Shot Prompting",
    italic=True,
    size=12,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=10,
)
add_paragraph(doc, f"Name: {NAME}")
add_paragraph(doc, f"USN: {USN}")
add_paragraph(doc, f"Date: {DATE_STR}", space_after=10)

add_heading(doc, "Introduction")
add_paragraph(
    doc,
    "This experiment studies how prompt structure changes the behavior of a large language model. Instead of using only a plain instruction, the lab practices several prompting strategies such as Chain of Thought, tabular formatting, fill-in-the-blank prompts, RGC framing, zero-shot prompting, one-shot prompting, and few-shot prompting.",
)
add_paragraph(
    doc,
    "The strongest evidence for this lab already exists in the local repository. The notebook GENAI_LAB3,4.ipynb contains saved prompt-response outputs for all the required prompting types, and the Actual Lab 3 folder contains additional prompt design material that explains advanced prompt patterns in more detail.",
)

add_heading(doc, "Objective")
for item in [
    "To practice multiple prompt engineering methods in a single workflow.",
    "To compare how prompt style affects structure, clarity, and accuracy of the output.",
    "To understand when zero-shot, one-shot, and few-shot prompting are useful.",
    "To build a small evidence-backed prompt library for future AI tasks.",
]:
    add_bullet(doc, item)

add_heading(doc, "Tools and Sources Used")
add_tools_table(
    doc,
    [
        ("GENAI_LAB3,4.ipynb", "Primary source of saved prompts and outputs for the required techniques."),
        ("prompting_examples.md", "Reference for richer prompt patterns including RGC, zero-shot, one-shot, and few-shot usage."),
        ("Lab 3 prompt.docx", "Example of prompt-driven response design using CoT, root prompt, and RGC framing."),
        ("Python notebook environment", "Execution environment used to run the prompting examples."),
    ],
)

add_heading(doc, "Methodology / Procedure")
for item in [
    "Review the local notebook and identify the cells corresponding to each prompting strategy.",
    "Extract the prompt text and the saved output already present in the notebook.",
    "Map the observed examples to the manual requirements for Chain of Thought, tabular format, fill-in-the-blank, RGC, zero-shot, one-shot, and few-shot prompting.",
    "Compare the simple notebook examples with the more advanced templates stored in the Actual Lab 3 folder.",
    "Summarize the strengths and limitations of each prompting style using the collected outputs.",
]:
    add_bullet(doc, item)

add_heading(doc, "Implementation Summary")
add_paragraph(
    doc,
    "The notebook first establishes the OpenAI-compatible client configuration and then executes a sequence of prompt experiments. The first group contains foundational patterns such as personal prompt, cognitive verifier, question refinement, and root prompt. The second group directly matches the lab requirement by showing Chain of Thought, tabular format, fill-in-the-blank, RGC prompting, zero-shot prompting, one-shot prompting, and few-shot prompting.",
)
add_paragraph(
    doc,
    "The supporting markdown and document files extend these ideas from short demonstrations into reusable templates. For example, prompting_examples.md shows how RGC can be embedded into a production-style system prompt, while Lab 3 prompt.docx demonstrates how multiple prompting ideas can be combined into a longer, stylistically controlled generation task.",
)

add_heading(doc, "Representative Prompts Used")
for line in [
    'Chain of Thought: {"role": "user", "content": "Solve step by step: A model has 80% accuracy on 50 samples. How many are correct?"}',
    'Tabular format: {"role": "user", "content": "Compare AI, ML, and DL in a table."}',
    'Fill-in-the-blank: {"role": "user", "content": "Fill in the blank: Machine Learning is a subset of ____."}',
    'RGC: {"role": "user", "content": "Role: Data Scientist Goal: Explain Overfitting Constraint: Use simple words and limit to 70 words"}',
    'Zero-shot: {"role": "user", "content": "Define Artificial Intelligence."}',
    'One-shot: {"role": "user", "content": "Dog => Animal\\nCat => ?"}',
    'Few-shot: {"role": "user", "content": "Dog => Animal\\nRose => Flower\\nApple => ?"}',
]:
    add_code_line(doc, line)
doc.add_paragraph()

add_heading(doc, "Observed Outputs from the Notebook")
add_paragraph(
    doc,
    "The following outputs were reviewed directly from the saved notebook results and are included here as evidence that the prompting patterns were already executed successfully.",
)
for item in [
    "Chain of Thought output: the model solved the accuracy problem step by step, identified 50 as the total number of samples, multiplied by 80%, and concluded that 40 predictions were correct.",
    "Tabular output: the model returned a structured comparison table with AI, ML, and DL as columns and features such as scope, data dependence, and complexity as comparison fields.",
    "Fill-in-the-blank output: the model answered that Machine Learning is a subset of Artificial Intelligence.",
    "RGC output: the model explained overfitting in simple language and within the requested word constraint, showing that role and constraint framing changed the style and length of the response.",
    "Zero-shot output: the model produced a direct definition of Artificial Intelligence without any example demonstrations being supplied.",
    "One-shot output: after seeing the relation 'Dog => Animal', the model completed 'Cat => Animal', showing format imitation from a single example.",
    "Few-shot output: after seeing multiple relation examples, the model completed 'Apple => Fruit', showing stronger pattern completion from more than one example.",
]:
    add_bullet(doc, item)

add_heading(doc, "Analysis of Prompting Techniques")
add_paragraph(
    doc,
    "Chain of Thought prompting was effective for procedural reasoning because it encouraged the model to show an ordered problem-solving path rather than only the final answer. This is helpful in educational settings where the logic matters as much as the result.",
)
add_paragraph(
    doc,
    "Tabular prompting improved readability and comparison quality. The output became easier to inspect because the model was forced to align concepts under a repeated structure instead of producing free-form prose.",
)
add_paragraph(
    doc,
    "Fill-in-the-blank prompting worked well for quick factual recall. In contrast, RGC prompting shaped not only the content but also the persona, task objective, and response constraint, which made the answer more controlled and context-aware.",
)
add_paragraph(
    doc,
    "Zero-shot prompting required the model to generalize from the instruction alone, while one-shot and few-shot prompting progressively improved pattern fidelity by showing example relationships first. The notebook outputs clearly demonstrate this transition from direct instruction to example-guided completion.",
)

add_heading(doc, "Supporting Evidence from Present Folder")
add_paragraph(
    doc,
    "The Actual Lab 3 folder contains additional evidence that the student practiced advanced prompt design beyond the short notebook cells. In particular, Lab 3 prompt.docx contains a long-form example that combines root prompt instructions, Chain of Thought reasoning, RGC framing, and stylistic control in a creative-response task. The prompting_examples.md file also contains reusable templates for root prompts, few-shot prompting, tabular outputs, fill-in-the-blank patterns, refinement patterns, and RGC-only prompts.",
)

add_heading(doc, "Observations")
for item in [
    "Prompt structure has a direct effect on output structure.",
    "Role and constraints improve control over tone, brevity, and task focus.",
    "Example-based prompting is especially useful when output format must be consistent.",
    "Saved notebook outputs were sufficient to replace most of the placeholder evidence in the first draft of the report.",
]:
    add_bullet(doc, item)

add_heading(doc, "Pending Evidence from Student")
add_paragraph(
    doc,
    "Most of the conceptual evidence for this lab is already present in the repository. The only optional additions that may still improve the final submission are presentation-oriented rather than content-critical.",
)
for item in [
    "Notebook screenshots for the prompt/output cells, if your faculty expects image evidence instead of text summary.",
    "Any ChatGPT or Gemini outputs you want added as extra comparison evidence beyond the saved notebook results.",
]:
    add_bullet(doc, item)

add_heading(doc, "Conclusion")
add_paragraph(
    doc,
    "This lab successfully demonstrates that prompting is a controllable design tool rather than a single generic instruction. The saved outputs in GENAI_LAB3,4.ipynb provide concrete evidence for Chain of Thought, tabular, fill-in-the-blank, RGC, zero-shot, one-shot, and few-shot prompting.",
)
add_paragraph(
    doc,
    "Overall, the review shows that the repository already contains meaningful lab 4 evidence. By incorporating those outputs into the document, the report becomes substantially stronger and much closer to a submission-ready lab record.",
)

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT_PATH))
print(OUTPUT_PATH)
